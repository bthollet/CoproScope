from __future__ import annotations

import shutil
import tempfile
import unittest
import zipfile
from pathlib import Path

from coproscope.core.common import RunContext, load_instance, read_csv, relative_to, write_csv
from coproscope.core.share import audit_repo
from coproscope.modules import biffageops, docuscope, privacyops


class PrivacyOpsTests(unittest.TestCase):
    def setUp(self) -> None:
        repo_root = Path(__file__).resolve().parents[2]
        source_example = repo_root / "examples" / "synthetic_copro"
        self.tempdir = tempfile.TemporaryDirectory()
        self.example_root = Path(self.tempdir.name) / "synthetic_copro"
        shutil.copytree(source_example, self.example_root)
        self.instance = load_instance(str(self.example_root / "instance.yml"), None)

    def tearDown(self) -> None:
        self.tempdir.cleanup()

    def _row_for_file(self, file_name: str) -> dict[str, str]:
        _, rows = read_csv(self.instance.register("documents"))
        for row in rows:
            if row.get("file_name") == file_name:
                return row
        raise AssertionError(f"Missing row for {file_name}")

    def _screening_row_for_file(self, file_name: str) -> dict[str, str]:
        _, rows = read_csv(privacyops.privacy_screening_path(self.instance))
        for row in rows:
            if row.get("file_name") == file_name:
                return row
        raise AssertionError(f"Missing screening row for {file_name}")

    def _inject_dirty_docx_parts(self, path: Path) -> None:
        temp_path = path.with_name(path.name + ".dirty.tmp")
        rels = (
            b'<Relationship Id="rPrivateComments" '
            b'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments" '
            b'Target="comments.xml"/>'
            b'<Relationship Id="rPrivateOle" '
            b'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/oleObject" '
            b'Target="embeddings/oleObject1.bin"/>'
            b'<Relationship Id="rPrivateLink" '
            b'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink" '
            b'Target="https://example.invalid/Jean-Dupont" TargetMode="External"/>'
        )
        content_types = (
            b'<Override PartName="/word/comments.xml" '
            b'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/>'
            b'<Override PartName="/word/footnotes.xml" '
            b'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/>'
            b'<Override PartName="/word/header1.xml" '
            b'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml"/>'
            b'<Override PartName="/docProps/custom.xml" '
            b'ContentType="application/vnd.openxmlformats-officedocument.custom-properties+xml"/>'
            b'<Override PartName="/word/embeddings/oleObject1.bin" '
            b'ContentType="application/vnd.openxmlformats-officedocument.oleObject"/>'
        )
        additions = {
            "word/comments.xml": (
                b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                b'<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                b'<w:comment w:id="0" w:author="Jean Dupont">'
                b'<w:p><w:r><w:t>jean.dupont@example.com</w:t></w:r></w:p>'
                b'</w:comment></w:comments>'
            ),
            "word/footnotes.xml": (
                b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                b'<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                b'<w:footnote w:id="1"><w:p><w:r><w:t>Note jean.dupont@example.com</w:t></w:r></w:p></w:footnote>'
                b'</w:footnotes>'
            ),
            "word/header1.xml": (
                b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                b'<w:hdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                b'<w:p><w:r><w:t>Monsieur </w:t></w:r><w:r><w:t>Jean Dupont</w:t></w:r></w:p>'
                b'</w:hdr>'
            ),
            "word/footer1.xml": (
                b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                b'<w:ftr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                b'<w:p><w:del><w:r><w:delText>jean.dupont@example.com</w:delText></w:r></w:del></w:p>'
                b'</w:ftr>'
            ),
            "docProps/custom.xml": (
                b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                b'<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/custom-properties">'
                b'<property name="private"><vt:lpwstr xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">'
                b'Jean Dupont jean.dupont@example.com</vt:lpwstr></property></Properties>'
            ),
            "customXml/private_item.xml": b"<private>Jean Dupont jean.dupont@example.com</private>",
            "word/embeddings/oleObject1.bin": b"Jean Dupont jean.dupont@example.com",
        }
        with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(temp_path, "w", compression=zipfile.ZIP_DEFLATED) as target:
            for info in source.infolist():
                data = source.read(info)
                if info.filename == "[Content_Types].xml":
                    data = data.replace(b"</Types>", content_types + b"</Types>")
                elif info.filename == "word/_rels/document.xml.rels":
                    data = data.replace(b"</Relationships>", rels + b"</Relationships>")
                target.writestr(info, data)
            for name, data in additions.items():
                target.writestr(name, data)
        temp_path.replace(path)

    def test_screen_existing_enriches_access_policy_for_already_deposited_docs(self) -> None:
        sensitive = self.example_root / "raw" / "2025_impayes_nominatifs.txt"
        sensitive.write_text(
            "Impayes nominatif: Monsieur Jean Dupont, lot A12, "
            "jean.dupont@example.com, IBAN FR76 3000 6000 0112 3456 7890 189.\n",
            encoding="utf-8",
        )

        run = RunContext(self.instance, "privacy screen")
        result = privacyops.screen_existing(self.instance, run, include_generated=False)

        self.assertEqual(result["status"], "ok")
        self.assertTrue(Path(str(result["screening_register"])).exists())
        self.assertTrue(Path(str(result["report"])).exists())
        row = self._row_for_file("2025_impayes_nominatifs.txt")
        self.assertEqual(row["raw_max_college"], "C8_Restreint_Critique")
        self.assertEqual(row["derivative_max_college"], "C2_Coproprietaires")
        self.assertIn("human_review", row["required_transformations"])
        self.assertEqual(row["privacy_review_status"], "A_ARBITRER")
        self.assertEqual(row["review_status_recommendation"], "DIFFUSABLE_APRES_AGREGATION")
        self.assertEqual(row["review_justification_required"], "YES")
        self.assertEqual(row["ai_processing_ceiling"], "no_ai")
        self.assertNotIn("jean.dupont@example.com", row["screening_signals"])
        self.assertNotIn("Jean Dupont", row["screening_signals"])
        self.assertIn("EMAIL:count=1", row["screening_signals"])

    def test_publication_path_cannot_leave_sensitive_document_as_raw_publication(self) -> None:
        public_dir = self.example_root / "outputs" / "publication"
        public_dir.mkdir(parents=True, exist_ok=True)
        sensitive = public_dir / "impayes_publics.txt"
        sensitive.write_text(
            "Impayés nominatifs: Monsieur Jean Dupont, lot A12, IBAN FR76 3000 6000 0112 3456 7890 189.\n",
            encoding="utf-8",
        )

        run = RunContext(self.instance, "privacy public conflict")
        privacyops.screen_existing(self.instance, run, include_generated=True)
        row = self._row_for_file("impayes_publics.txt")

        self.assertEqual(row["raw_max_college"], "C8_Restreint_Critique")
        self.assertEqual(row["derivative_max_college"], "C2_Coproprietaires")
        self.assertNotEqual(row["publication_form"], "raw")
        self.assertIn("aggregation", row["required_transformations"])
        self.assertIn("human_review", row["required_transformations"])
        self.assertEqual(row["privacy_review_status"], "BLOQUE")
        self.assertIn("chemin de diffusion large", row["publication_blocker"])

    def test_clear_ag_path_can_broaden_default_c4_to_c2(self) -> None:
        ag_dir = self.example_root / "220_Assemblees_generales" / "221_AG_par_date" / "2026-04-29_AGE"
        ag_dir.mkdir(parents=True, exist_ok=True)
        source = ag_dir / "2026-04-29_convocation.md"
        source.write_text("Convocation assemblee generale sans donnees nominatives.\n", encoding="utf-8")

        run = RunContext(self.instance, "privacy ag path")
        privacyops.screen_existing(self.instance, run, include_generated=False, scan_workspace_prefixes=True)
        row = self._row_for_file("2026-04-29_convocation.md")

        self.assertEqual(row["raw_max_college"], "C2_Coproprietaires")
        self.assertEqual(row["derivative_max_college"], "C2_Coproprietaires")
        self.assertEqual(row["publication_form"], "raw")
        self.assertEqual(row["privacy_review_status"], "DIFFUSABLE_BRUT")

    def test_human_review_decision_requires_justification_for_wide_diffusion(self) -> None:
        ag_dir = self.example_root / "220_Assemblees_generales" / "221_AG_par_date" / "2026-04-29_AGE"
        ag_dir.mkdir(parents=True, exist_ok=True)
        source = ag_dir / "2026-04-29_convocation.md"
        source.write_text("Convocation assemblee generale sans donnees nominatives.\n", encoding="utf-8")

        run = RunContext(self.instance, "privacy human review")
        privacyops.screen_existing(self.instance, run, include_generated=False, scan_workspace_prefixes=True)
        row = self._row_for_file("2026-04-29_convocation.md")

        with self.assertRaises(ValueError):
            privacyops.record_human_review_decision(
                self.instance,
                run,
                doc_id=row["doc_id"],
                status="DIFFUSABLE_BRUT",
            )

        result = privacyops.record_human_review_decision(
            self.instance,
            run,
            doc_id=row["doc_id"],
            status="DIFFUSABLE_BRUT",
            justification="Convocation AG sans donnees nominatives ni signal sensible.",
            reviewer="CS",
        )

        self.assertEqual(result["status"], "ok")
        reviewed = self._row_for_file("2026-04-29_convocation.md")
        screened = self._screening_row_for_file("2026-04-29_convocation.md")
        self.assertEqual(reviewed["privacy_review_status"], "DIFFUSABLE_BRUT")
        self.assertEqual(screened["privacy_review_status"], "DIFFUSABLE_BRUT")
        self.assertEqual(reviewed["privacy_review_owner"], "CS")
        self.assertIn("Convocation AG", reviewed["privacy_review_justification"])

    def test_human_review_cannot_mark_required_transformation_as_raw_diffusion(self) -> None:
        source = self.example_root / "raw" / "note_a_biffer_brut.md"
        source.write_text(
            "Note de travail avec Monsieur Jean Dupont et jean.dupont@example.com.\n",
            encoding="utf-8",
        )

        run = RunContext(self.instance, "privacy raw forbidden")
        privacyops.screen_existing(self.instance, run, include_generated=False)
        row = self._row_for_file("note_a_biffer_brut.md")

        with self.assertRaises(ValueError):
            privacyops.record_human_review_decision(
                self.instance,
                run,
                doc_id=row["doc_id"],
                status="DIFFUSABLE_BRUT",
                justification="Decision humaine explicite.",
                reviewer="CS",
            )

    def test_path_rules_do_not_scan_body_text_as_path_markers(self) -> None:
        ag_dir = self.example_root / "220_Assemblees_generales" / "221_AG_par_date" / "2026-04-29_AGE"
        ag_dir.mkdir(parents=True, exist_ok=True)
        source = ag_dir / "2026-04-29_note_finance_generique.md"
        source.write_text(
            "La convocation mentionne la banque et les factures comme sujets generiques sans identifiant direct.\n",
            encoding="utf-8",
        )

        run = RunContext(self.instance, "privacy path body")
        privacyops.screen_existing(self.instance, run, include_generated=False, scan_workspace_prefixes=True)
        row = self._row_for_file("2026-04-29_note_finance_generique.md")

        self.assertEqual(row["raw_max_college"], "C4_Conseil_Syndical")
        self.assertEqual(row["derivative_max_college"], "C2_Coproprietaires")
        self.assertEqual(row["publication_form"], "redaction_required")

    def test_security_secret_redaction_removes_secret_value(self) -> None:
        source = self.example_root / "raw" / "codes_badges.md"
        source.write_text("Code badge: 48291 et mot de passe: AZERTY-42\n", encoding="utf-8")

        run = RunContext(self.instance, "privacy secret redact")
        docuscope.inventory(self.instance, run)
        privacyops.screen_existing(self.instance, run, include_generated=False)
        row = self._row_for_file("codes_badges.md")

        result = biffageops.redact_document(
            self.instance,
            run,
            doc_id=row["doc_id"],
            mode="redaction_irreversible",
            target_college="C2_Coproprietaires",
        )

        content = Path(str(result["redacted_path"])).read_text(encoding="utf-8")
        self.assertNotIn("48291", content)
        self.assertNotIn("AZERTY-42", content)
        self.assertIn("[BIFFE_SECURITY_SECRET]", content)

    def test_screen_existing_uses_existing_ocr_text_when_native_sample_is_empty(self) -> None:
        source = self.example_root / "raw" / "scan_impayes.png"
        source.write_bytes(b"not a real image")

        run = RunContext(self.instance, "privacy ocr text")
        docuscope.inventory(self.instance, run)
        fields, rows = read_csv(self.instance.register("documents"))
        row = self._row_for_file("scan_impayes.png")
        text_path = self.example_root / "staging" / "text" / f"{row['doc_id']}.ocr.txt"
        text_path.parent.mkdir(parents=True, exist_ok=True)
        text_path.write_text("Impayés: Monsieur Jean Dupont, IBAN FR76 3000 6000 0112 3456 7890 189.\n", encoding="utf-8")
        for item in rows:
            if item.get("doc_id") == row["doc_id"]:
                item["text_path"] = relative_to(self.example_root, text_path)
                item["status_ocr"] = "OCR_DONE"
        write_csv(self.instance.register("documents"), fields, rows)

        privacyops.screen_existing(self.instance, run, include_generated=False)
        screened = self._row_for_file("scan_impayes.png")

        self.assertEqual(screened["raw_max_college"], "C8_Restreint_Critique")
        self.assertEqual(screened["ai_processing_ceiling"], "no_ai")

    def test_biffageops_pseudonymizes_text_and_keeps_mapping_restricted(self) -> None:
        source = self.example_root / "raw" / "note_a_biffer.md"
        source.write_text(
            "Note diffusable apres biffage. Monsieur Jean Dupont joignable a jean.dupont@example.com.\n",
            encoding="utf-8",
        )
        run = RunContext(self.instance, "privacy redact")
        docuscope.inventory(self.instance, run)
        docuscope.extract_text(self.instance, run)
        privacyops.screen_existing(self.instance, run, include_generated=False)
        row = self._row_for_file("note_a_biffer.md")

        result = biffageops.redact_document(
            self.instance,
            run,
            doc_id=row["doc_id"],
            mode="pseudonymisation_tracee",
            target_college="C2_Coproprietaires",
        )

        self.assertEqual(result["status"], "ok")
        redacted = Path(str(result["redacted_path"]))
        self.assertTrue(redacted.exists())
        content = redacted.read_text(encoding="utf-8")
        self.assertNotIn("jean.dupont@example.com", content)
        self.assertNotIn("Monsieur Jean Dupont", content)
        self.assertIn("EMAIL_", content)
        map_path = self.example_root / str(result["map_path"])
        self.assertTrue(map_path.exists())
        self.assertIn("C8_table_correspondance_biffage", str(map_path))

    def test_redaction_map_register_outside_restricted_area_is_not_used(self) -> None:
        self.instance.payload["registers"]["redaction_map"] = "./registers/table_correspondance_biffage.csv"
        source = self.example_root / "raw" / "note_map_location.md"
        source.write_text("Monsieur Jean Dupont jean.dupont@example.com\n", encoding="utf-8")

        run = RunContext(self.instance, "privacy map location")
        docuscope.inventory(self.instance, run)
        privacyops.screen_existing(self.instance, run, include_generated=False)
        row = self._row_for_file("note_map_location.md")

        result = biffageops.redact_document(
            self.instance,
            run,
            doc_id=row["doc_id"],
            mode="pseudonymisation_tracee",
            target_college="C2_Coproprietaires",
        )

        self.assertFalse((self.example_root / "registers" / "table_correspondance_biffage.csv").exists())
        self.assertIn("C8_table_correspondance_biffage", str(result["map_path"]))

    def test_redaction_rejects_target_wider_than_derivative_ceiling(self) -> None:
        source = self.example_root / "raw" / "note_target.md"
        source.write_text("Monsieur Jean Dupont jean.dupont@example.com\n", encoding="utf-8")

        run = RunContext(self.instance, "privacy target ceiling")
        docuscope.inventory(self.instance, run)
        privacyops.screen_existing(self.instance, run, include_generated=False)
        row = self._row_for_file("note_target.md")

        with self.assertRaises(ValueError):
            biffageops.redact_document(
                self.instance,
                run,
                doc_id=row["doc_id"],
                mode="redaction_irreversible",
                target_college="C0_Public",
            )

    def test_pdf_redaction_removes_underlying_text_when_pymupdf_is_available(self) -> None:
        try:
            import fitz  # type: ignore
        except ImportError:
            self.skipTest("PyMuPDF unavailable")

        source = self.example_root / "raw" / "note_pdf_a_biffer.pdf"
        document = fitz.open()
        page = document.new_page()
        page.insert_text((72, 72), "Monsieur Jean Dupont jean.dupont@example.com")
        document.save(str(source))
        document.close()

        run = RunContext(self.instance, "privacy pdf redact")
        docuscope.inventory(self.instance, run)
        privacyops.screen_existing(self.instance, run, include_generated=False)
        row = self._row_for_file("note_pdf_a_biffer.pdf")

        result = biffageops.redact_document(
            self.instance,
            run,
            doc_id=row["doc_id"],
            mode="redaction_irreversible",
            target_college="C2_Coproprietaires",
        )

        out_doc = fitz.open(str(result["redacted_path"]))
        try:
            text = "\n".join(out_doc.load_page(index).get_text("text") for index in range(out_doc.page_count))
        finally:
            out_doc.close()
        self.assertNotIn("jean.dupont@example.com", text)
        self.assertNotIn("Monsieur Jean Dupont", text)

        queue = biffageops.build_redaction_queue(self.instance, run)
        _, queue_rows = read_csv(Path(str(queue["queue"])))
        queued = next(item for item in queue_rows if item["doc_id"] == row["doc_id"])
        self.assertEqual(queued["queue_status"], "REDACTED")

    def test_pdf_redaction_can_fall_back_to_registered_text_derivative(self) -> None:
        source = self.example_root / "raw" / "scan_demande_a_biffer.pdf"
        source.write_bytes(b"not a real pdf")

        run = RunContext(self.instance, "privacy pdf text derivative")
        docuscope.inventory(self.instance, run)
        fields, rows = read_csv(self.instance.register("documents"))
        row = self._row_for_file("scan_demande_a_biffer.pdf")
        text_path = self.example_root / "staging" / "text" / f"{row['doc_id']}.ocr.txt"
        text_path.parent.mkdir(parents=True, exist_ok=True)
        text_path.write_text(
            "Demande nominative: Monsieur Jean Dupont, lot A12, jean.dupont@example.com.\n",
            encoding="utf-8",
        )
        for item in rows:
            if item.get("doc_id") == row["doc_id"]:
                item["text_path"] = relative_to(self.example_root, text_path)
                item["status_ocr"] = "OCR_DONE"
        write_csv(self.instance.register("documents"), fields, rows)
        privacyops.screen_existing(self.instance, run, include_generated=False)
        row = self._row_for_file("scan_demande_a_biffer.pdf")

        result = biffageops.redact_document(
            self.instance,
            run,
            doc_id=row["doc_id"],
            mode="redaction_irreversible",
            target_college="C2_Coproprietaires",
        )

        redacted = Path(str(result["redacted_path"]))
        self.assertEqual(redacted.suffix, ".txt")
        content = redacted.read_text(encoding="utf-8")
        self.assertNotIn("Jean Dupont", content)
        self.assertNotIn("jean.dupont@example.com", content)
        self.assertIn("[BIFFE_PERSON_NAME]", content)
        self.assertIn("[BIFFE_EMAIL]", content)

    def test_redaction_required_without_identifiers_registers_text_derivative(self) -> None:
        source = self.example_root / "raw" / "carte_professionnelle.pdf"
        source.write_bytes(b"not a real pdf")

        run = RunContext(self.instance, "privacy no identifiers derivative")
        docuscope.inventory(self.instance, run)
        fields, rows = read_csv(self.instance.register("documents"))
        row = self._row_for_file("carte_professionnelle.pdf")
        text_path = self.example_root / "staging" / "text" / f"{row['doc_id']}.native.txt"
        text_path.parent.mkdir(parents=True, exist_ok=True)
        text_path.write_text("Carte professionnelle valide jusqu'au 23/03/2026.\n", encoding="utf-8")
        for item in rows:
            if item.get("doc_id") == row["doc_id"]:
                item["text_path"] = relative_to(self.example_root, text_path)
                item["status_ocr"] = "TEXT_EXTRACTED"
                item["publication_form"] = "redaction_required"
                item["personal_data_level"] = "none"
        write_csv(self.instance.register("documents"), fields, rows)

        result = biffageops.redact_document(
            self.instance,
            run,
            doc_id=row["doc_id"],
            mode="redaction_irreversible",
            target_college="C2_Coproprietaires",
        )

        self.assertEqual(result["status"], "ok")
        self.assertEqual(result["match_count"], 0)
        redacted = Path(str(result["redacted_path"]))
        self.assertEqual(redacted.suffix, ".txt")
        self.assertIn("Carte professionnelle", redacted.read_text(encoding="utf-8"))
        refreshed = self._row_for_file("carte_professionnelle.pdf")
        self.assertEqual(refreshed["redaction_status"], "REDACTED")

    def test_docx_redaction_sanitizes_package_parts_and_hidden_values(self) -> None:
        try:
            from docx import Document  # type: ignore
        except ImportError:
            self.skipTest("python-docx unavailable")

        source = self.example_root / "raw" / "docx_sale.docx"
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.add_run("Monsieur ")
        paragraph.add_run("Jean ")
        paragraph.add_run("Dupont")
        document.add_paragraph("Contact jean.dupont@example.com")
        document.core_properties.author = "Jean Dupont"
        document.core_properties.comments = "jean.dupont@example.com"
        document.save(str(source))
        self._inject_dirty_docx_parts(source)

        run = RunContext(self.instance, "privacy docx sanitize")
        docuscope.inventory(self.instance, run)
        privacyops.screen_existing(self.instance, run, include_generated=False)
        row = self._row_for_file("docx_sale.docx")

        result = biffageops.redact_document(
            self.instance,
            run,
            doc_id=row["doc_id"],
            mode="redaction_irreversible",
            target_college="C2_Coproprietaires",
        )

        with zipfile.ZipFile(str(result["redacted_path"])) as package:
            names = package.namelist()
            self.assertFalse(any(name.startswith("docProps/") for name in names))
            self.assertFalse(any(name.startswith("customXml/") for name in names))
            self.assertFalse(any(name.startswith("word/embeddings/") for name in names))
            self.assertNotIn("word/comments.xml", names)
            rels = package.read("word/_rels/document.xml.rels").decode("utf-8")
            self.assertNotIn("comments.xml", rels)
            self.assertNotIn("oleObject", rels)
            self.assertNotIn("TargetMode=\"External\"", rels)
            package_bytes = b"\n".join(package.read(name) for name in names if not name.endswith("/"))

        self.assertNotIn(b"jean.dupont@example.com", package_bytes)
        self.assertNotIn(b"Jean Dupont", package_bytes)
        self.assertIn(b"[BIFFE_EMAIL]", package_bytes)
        reopened = Document(str(result["redacted_path"]))
        reopened_text = "\n".join(paragraph.text for paragraph in reopened.paragraphs)
        self.assertIn("[BIFFE_EMAIL]", reopened_text)

    def test_share_audit_blocks_redaction_correspondence_tables(self) -> None:
        repo_root = Path(__file__).resolve().parents[2]
        config_path = repo_root / "server" / "src" / "coproscope" / "configs" / "github_sharing.default.yml"
        with tempfile.TemporaryDirectory() as tmpdir:
            audit_root = Path(tmpdir)
            target = audit_root / "examples" / "synthetic_copro" / "raw"
            target.mkdir(parents=True)
            (target / "table_correspondance_biffage.csv").write_text("secret\n", encoding="utf-8")
            (audit_root / "README.md").write_text("# repo\n", encoding="utf-8")

            result = audit_repo(audit_root, config_path)

            self.assertIn("README.md", result["shareable"])
            self.assertIn("examples/synthetic_copro/raw/table_correspondance_biffage.csv", result["blocked"])


if __name__ == "__main__":
    unittest.main()
