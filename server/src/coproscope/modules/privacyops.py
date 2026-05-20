from __future__ import annotations

from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable

from ..core.common import (
    DEFAULT_DOCUMENT_FIELDS,
    InstanceConfig,
    RunContext,
    now_iso,
    read_csv,
    relative_to,
    sha256_file,
    write_csv,
    write_text,
)
from ..core.privacy import (
    SCREENING_FIELDS,
    apply_access_policy,
    ensure_policy_fields,
    exposure_risk_for,
)
from ..core.sanitizer import extract_docx_package_text


SKIP_DIR_NAMES = {
    ".git",
    ".venv",
    "__pycache__",
    ".pytest_cache",
    ".mypy_cache",
    ".ruff_cache",
    ".secrets",
    "node_modules",
    "models",
    "huggingface",
    "privacy_dir",
    "redacted_dir",
    "dossier_confidentialite",
    "dossier_biffages",
    "biffageops",
}
SKIP_FILE_NAMES = {"desktop.ini", "thumbs.db"}
SKIP_FILE_MARKERS = {
    "table_correspondance",
    "c8_table_correspondance_biffage",
    "registre_screening_confidentialite",
    "registre_biffages",
    "file_biffage",
}
DEFAULT_WORKSPACE_PREFIXES = (
    "100_",
    "200_",
    "210_",
    "220_",
    "230_",
    "240_",
    "250_",
    "260_",
    "270_",
    "280_",
    "290_",
)
TEXT_EXTENSIONS = {"txt", "md", "csv", "tsv", "json", "yml", "yaml", "html", "htm"}


def _privacy_dir(instance: InstanceConfig) -> Path:
    path = instance.artifact("privacy_dir")
    path.mkdir(parents=True, exist_ok=True)
    return path


def privacy_screening_path(instance: InstanceConfig) -> Path:
    try:
        return instance.register("privacy_screening")
    except KeyError:
        return _privacy_dir(instance) / "registre_screening_confidentialite.csv"


def redactions_path(instance: InstanceConfig) -> Path:
    try:
        return instance.register("redactions")
    except KeyError:
        return _privacy_dir(instance) / "registre_biffages.csv"


def _read_text_file(path: Path, max_chars: int) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="ignore")[:max_chars]
    except OSError:
        return ""


def _read_pdf_text(path: Path, max_chars: int) -> str:
    try:
        import fitz  # type: ignore

        document = fitz.open(str(path))
        try:
            parts: list[str] = []
            for index in range(min(document.page_count, 10)):
                parts.append(document.load_page(index).get_text("text"))
                if sum(len(part) for part in parts) >= max_chars:
                    break
            return "\n".join(parts)[:max_chars]
        finally:
            document.close()
    except Exception:  # noqa: BLE001
        return ""


def _read_docx_text(path: Path, max_chars: int) -> str:
    parts: list[str] = []
    try:
        from docx import Document  # type: ignore

        document = Document(str(path))
        parts.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
        for table in document.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if values:
                    parts.append(" | ".join(values))
    except Exception:  # noqa: BLE001
        pass
    try:
        parts.append(extract_docx_package_text(path, max_chars=max_chars))
    except Exception:  # noqa: BLE001
        pass
    return "\n".join(part for part in parts if part)[:max_chars]


def _read_xlsx_text(path: Path, max_chars: int) -> str:
    try:
        from openpyxl import load_workbook  # type: ignore

        workbook = load_workbook(str(path), read_only=True, data_only=True)
        try:
            parts: list[str] = []
            for worksheet in workbook.worksheets:
                parts.append(f"FEUILLE {worksheet.title}")
                for row in worksheet.iter_rows(values_only=True):
                    values = [str(value).strip() if value is not None else "" for value in row]
                    if any(values):
                        parts.append(" | ".join(values))
                    if sum(len(part) for part in parts) >= max_chars:
                        return "\n".join(parts)[:max_chars]
        finally:
            workbook.close()
        return "\n".join(parts)[:max_chars]
    except Exception:  # noqa: BLE001
        return ""


def _read_sample(path: Path, max_chars: int) -> str:
    ext = path.suffix.lower().lstrip(".")
    if ext in TEXT_EXTENSIONS:
        return _read_text_file(path, max_chars)
    if ext == "pdf":
        return _read_pdf_text(path, max_chars)
    if ext == "docx":
        return _read_docx_text(path, max_chars)
    if ext in {"xlsx", "xlsm"}:
        return _read_xlsx_text(path, max_chars)
    return ""


def _read_existing_text_artifact(instance: InstanceConfig, row: dict[str, str], max_chars: int) -> str:
    text_path = row.get("text_path", "")
    if not text_path:
        return ""
    path = Path(text_path)
    if not path.is_absolute():
        path = instance.root("workspace") / path
    if not path.exists() or not path.is_file():
        return ""
    return _read_text_file(path, max_chars)


def _scan_roots(instance: InstanceConfig, include_generated: bool) -> list[Path]:
    roots: list[Path] = []
    workspace = instance.root("workspace")
    for root_name in ("raw", "system"):
        try:
            roots.append(instance.root(root_name))
        except KeyError:
            pass
    if include_generated:
        for root_name in ("outputs", "staging"):
            try:
                roots.append(instance.root(root_name))
            except KeyError:
                pass
    roots.extend(instance.root_list("restricted"))

    for child in workspace.iterdir() if workspace.exists() else []:
        if child.is_dir() and child.name.startswith(DEFAULT_WORKSPACE_PREFIXES):
            roots.append(child)

    unique: list[Path] = []
    seen: set[str] = set()
    for root in roots:
        resolved = root.resolve()
        key = str(resolved).lower()
        if resolved.exists() and key not in seen:
            seen.add(key)
            unique.append(resolved)
    return unique


def _iter_files(roots: Iterable[Path]):
    for root in roots:
        for path in sorted(root.rglob("*")):
            if not path.is_file():
                continue
            if path.name.lower() in SKIP_FILE_NAMES:
                continue
            if any(marker in path.name.lower() for marker in SKIP_FILE_MARKERS):
                continue
            if any(part.lower() in SKIP_DIR_NAMES for part in path.parts):
                continue
            yield path


def _source_zone(workspace_root: Path, path: Path) -> str:
    try:
        rel = path.resolve().relative_to(workspace_root.resolve())
        return rel.parts[0] if rel.parts else "workspace"
    except ValueError:
        return path.parent.name or "external"


def _doc_id_for(digest: str) -> str:
    return f"DOC-{digest[:12].upper()}"


def _existing_indexes(rows: list[dict[str, str]]) -> tuple[dict[str, dict[str, str]], dict[str, dict[str, str]]]:
    by_path = {row.get("original_path", "").replace("\\", "/").lower(): row for row in rows if row.get("original_path")}
    by_digest = {row.get("sha256", ""): row for row in rows if row.get("sha256")}
    return by_path, by_digest


def _base_document_row(
    instance: InstanceConfig,
    workspace_root: Path,
    path: Path,
    digest: str,
    previous: dict[str, str],
    seen_at: str,
) -> dict[str, str]:
    stat = path.stat()
    relative_path = relative_to(workspace_root, path)
    row = {field: "" for field in DEFAULT_DOCUMENT_FIELDS}
    row.update(previous)
    row.update(
        {
            "doc_id": previous.get("doc_id") or _doc_id_for(digest),
            "instance_id": instance.instance_id,
            "entity_id": instance.entity_id,
            "scope": previous.get("scope") or "screening",
            "sha256": digest,
            "original_path": relative_path,
            "file_name": path.name,
            "extension": path.suffix.lower().lstrip("."),
            "size_bytes": str(stat.st_size),
            "last_modified": datetime.fromtimestamp(stat.st_mtime, timezone.utc).replace(microsecond=0).isoformat(),
            "first_seen": previous.get("first_seen") or seen_at,
            "source_zone": previous.get("source_zone") or _source_zone(workspace_root, path),
            "source_kind": previous.get("source_kind") or "screened_existing",
            "status_ocr": previous.get("status_ocr") or "PENDING",
            "classification_status": previous.get("classification_status") or "PENDING",
        }
    )
    return row


def _screening_row(row: dict[str, str]) -> dict[str, str]:
    priority, exposure_risk, action = exposure_risk_for(row)
    values = {field: row.get(field, "") for field in SCREENING_FIELDS}
    values.update(
        {
            "exposure_risk": exposure_risk,
            "remediation_priority": priority,
            "recommended_action": action,
        }
    )
    return values


def _write_report(instance: InstanceConfig, screening_rows: list[dict[str, str]]) -> Path:
    report_path = instance.artifact("reports_dir") / "rapport_screening_confidentialite.md"
    by_raw = Counter(row.get("raw_max_college", "") for row in screening_rows)
    by_priority = Counter(row.get("remediation_priority", "") for row in screening_rows)
    to_redact = sum(1 for row in screening_rows if "redaction" in row.get("required_transformations", ""))
    review = sum(1 for row in screening_rows if row.get("review_required"))

    lines = [
        "# Rapport de screening confidentialite",
        "",
        f"- Instance: {instance.display_name}",
        f"- Documents screenes: {len(screening_rows)}",
        f"- Documents a biffer: {to_redact}",
        f"- Documents a revoir: {review}",
        "",
        "## Repartition par college brut",
        "",
        "| College | Nombre |",
        "| --- | ---: |",
    ]
    for college, count in sorted(by_raw.items()):
        lines.append(f"| {college or '-'} | {count} |")
    lines.extend(
        [
            "",
            "## Priorites de remediation",
            "",
            "| Priorite | Nombre |",
            "| --- | ---: |",
        ]
    )
    for priority, count in sorted(by_priority.items()):
        lines.append(f"| {priority or '-'} | {count} |")
    lines.extend(
        [
            "",
            "## Points a traiter",
            "",
            "| Priorite | Risque | Document | College brut | Transformation | Action |",
            "| --- | --- | --- | --- | --- | --- |",
        ]
    )
    priority_rows = [
        row
        for row in screening_rows
        if row.get("remediation_priority") in {"P0", "P1", "P2"}
    ]
    for row in priority_rows[:50]:
        lines.append(
            " | ".join(
                [
                    "",
                    row.get("remediation_priority", ""),
                    row.get("exposure_risk", ""),
                    row.get("doc_id", ""),
                    row.get("raw_max_college", ""),
                    row.get("required_transformations", ""),
                    row.get("recommended_action", ""),
                    "",
                ]
            )
        )
    if not priority_rows:
        lines.append("| OK | NO_OBVIOUS_EXPOSURE | - | - | - | Aucun point prioritaire |")
    write_text(report_path, "\n".join(lines) + "\n")
    return report_path


def screen_existing(
    instance: InstanceConfig,
    run: RunContext,
    *,
    include_generated: bool = True,
    max_text_chars: int = 50000,
) -> dict[str, object]:
    workspace_root = instance.root("workspace")
    documents_path = instance.register("documents")
    fields, document_rows = read_csv(documents_path)
    if not fields:
        fields = list(DEFAULT_DOCUMENT_FIELDS)
    for field in DEFAULT_DOCUMENT_FIELDS:
        if field not in fields:
            fields.append(field)
    ensure_policy_fields(fields)
    by_path, by_digest = _existing_indexes(document_rows)
    seen_at = now_iso()
    updated_by_path: dict[str, dict[str, str]] = {
        row.get("original_path", "").replace("\\", "/").lower(): row for row in document_rows if row.get("original_path")
    }
    screening_rows: list[dict[str, str]] = []

    for path in _iter_files(_scan_roots(instance, include_generated)):
        try:
            digest = sha256_file(path)
            relative_path = relative_to(workspace_root, path)
            key = relative_path.replace("\\", "/").lower()
            previous = by_path.get(key) or by_digest.get(digest, {})
            row = _base_document_row(instance, workspace_root, path, digest, previous, seen_at)
        except OSError as exc:
            run.log_error(f"privacy screen skipped unreadable file {path}: {exc}")
            continue
        text = _read_sample(path, max_text_chars)
        if not text.strip():
            text = _read_existing_text_artifact(instance, row, max_text_chars)
        apply_access_policy(row, text=text, instance=instance)
        updated_by_path[key] = row
        screening_rows.append(_screening_row(row))

    document_rows = sorted(updated_by_path.values(), key=lambda item: (item.get("original_path", "").lower(), item.get("doc_id", "")))
    write_csv(documents_path, fields, document_rows)
    screening_path = privacy_screening_path(instance)
    write_csv(screening_path, SCREENING_FIELDS, screening_rows)
    report_path = _write_report(instance, screening_rows)
    run.log_action("privacy_screen_existing", screening_path, f"rows={len(screening_rows)}")
    run.log_action("write", documents_path, f"privacy enriched documents={len(document_rows)}")
    run.log_action("write", report_path, "privacy screening report")
    return {
        "status": "ok",
        "screened_count": len(screening_rows),
        "document_count": len(document_rows),
        "screening_register": str(screening_path),
        "report": str(report_path),
        "by_raw_college": dict(Counter(row.get("raw_max_college", "") for row in screening_rows)),
        "priority_counts": dict(Counter(row.get("remediation_priority", "") for row in screening_rows)),
    }


def screening_summary(instance: InstanceConfig) -> dict[str, object]:
    path = privacy_screening_path(instance)
    _, rows = read_csv(path)
    return {
        "path": str(path),
        "count": len(rows),
        "by_raw_college": dict(Counter(row.get("raw_max_college", "") for row in rows)),
        "priority_counts": dict(Counter(row.get("remediation_priority", "") for row in rows)),
    }
