from __future__ import annotations

from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable

from ..core.common import (
    DEFAULT_DOCUMENT_FIELDS,
    InstanceConfig,
    RunContext,
    now_iso,
    read_csv,
    relative_to,
    sha256_file,
    write_csv,
    write_text,
)
from ..core.privacy import (
    ACCESS_RANK,
    SCREENING_FIELDS,
    apply_access_policy,
    ensure_policy_fields,
    exposure_risk_for,
    normalize_college,
)
from ..core.sanitizer import extract_docx_package_text


SKIP_DIR_NAMES = {
    ".git",
    ".venv",
    "__pycache__",
    ".pytest_cache",
    ".mypy_cache",
    ".ruff_cache",
    ".secrets",
    "node_modules",
    "coproscope",
    "server",
    "docs",
    "_archives",
    "_instance_docs",
    "_transition_reports",
    "_worktrees",
    "900_systeme_audit",
    "990_archives",
    "models",
    "huggingface",
    "privacy_dir",
    "redacted_dir",
    "dossier_confidentialite",
    "dossier_biffages",
    "biffageops",
}
SKIP_FILE_NAMES = {"desktop.ini", "thumbs.db", "passation_coproscope_local.md"}
SKIP_FILE_MARKERS = {
    "table_correspondance",
    "c8_table_correspondance_biffage",
    "registre_screening_confidentialite",
    "registre_biffages",
    "file_biffage",
}
DEFAULT_WORKSPACE_PREFIXES = (
    "100_",
    "200_",
    "210_",
    "220_",
    "230_",
    "240_",
    "250_",
    "260_",
    "270_",
    "280_",
    "290_",
)
TEXT_EXTENSIONS = {"txt", "md", "csv", "tsv", "json", "yml", "yaml", "html", "htm"}
REVIEW_STATUSES = {
    "DIFFUSABLE_BRUT",
    "DIFFUSABLE_APRES_BIFFAGE",
    "DIFFUSABLE_APRES_AGREGATION",
    "BLOQUE",
    "A_ARBITRER",
}
DIFFUSION_REVIEW_STATUSES = {
    "DIFFUSABLE_BRUT",
    "DIFFUSABLE_APRES_BIFFAGE",
    "DIFFUSABLE_APRES_AGREGATION",
}
REVIEW_DECISION_FIELDS = [
    "review_status_recommendation",
    "review_next_step",
    "review_justification_required",
    "privacy_review_justification",
    "privacy_review_owner",
    "privacy_reviewed_at",
    "publication_blocker",
]
SCREENING_REVIEW_FIELDS = list(SCREENING_FIELDS)
for _field in ["privacy_review_status", *REVIEW_DECISION_FIELDS]:
    if _field not in SCREENING_REVIEW_FIELDS:
        SCREENING_REVIEW_FIELDS.append(_field)


def _privacy_dir(instance: InstanceConfig) -> Path:
    path = instance.artifact("privacy_dir")
    path.mkdir(parents=True, exist_ok=True)
    return path


def privacy_screening_path(instance: InstanceConfig) -> Path:
    try:
        return instance.register("privacy_screening")
    except KeyError:
        return _privacy_dir(instance) / "registre_screening_confidentialite.csv"


def redactions_path(instance: InstanceConfig) -> Path:
    try:
        return instance.register("redactions")
    except KeyError:
        return _privacy_dir(instance) / "registre_biffages.csv"


def _read_text_file(path: Path, max_chars: int) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="ignore")[:max_chars]
    except OSError:
        return ""


def _read_pdf_text(path: Path, max_chars: int) -> str:
    try:
        import fitz  # type: ignore

        document = fitz.open(str(path))
        try:
            parts: list[str] = []
            for index in range(min(document.page_count, 10)):
                parts.append(document.load_page(index).get_text("text"))
                if sum(len(part) for part in parts) >= max_chars:
                    break
            return "\n".join(parts)[:max_chars]
        finally:
            document.close()
    except Exception:  # noqa: BLE001
        return ""


def _read_docx_text(path: Path, max_chars: int) -> str:
    parts: list[str] = []
    try:
        from docx import Document  # type: ignore

        document = Document(str(path))
        parts.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
        for table in document.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if values:
                    parts.append(" | ".join(values))
    except Exception:  # noqa: BLE001
        pass
    try:
        parts.append(extract_docx_package_text(path, max_chars=max_chars))
    except Exception:  # noqa: BLE001
        pass
    return "\n".join(part for part in parts if part)[:max_chars]


def _read_xlsx_text(path: Path, max_chars: int) -> str:
    try:
        from openpyxl import load_workbook  # type: ignore

        workbook = load_workbook(str(path), read_only=True, data_only=True)
        try:
            parts: list[str] = []
            for worksheet in workbook.worksheets:
                parts.append(f"FEUILLE {worksheet.title}")
                for row in worksheet.iter_rows(values_only=True):
                    values = [str(value).strip() if value is not None else "" for value in row]
                    if any(values):
                        parts.append(" | ".join(values))
                    if sum(len(part) for part in parts) >= max_chars:
                        return "\n".join(parts)[:max_chars]
        finally:
            workbook.close()
        return "\n".join(parts)[:max_chars]
    except Exception:  # noqa: BLE001
        return ""


def _read_sample(path: Path, max_chars: int) -> str:
    ext = path.suffix.lower().lstrip(".")
    if ext in TEXT_EXTENSIONS:
        return _read_text_file(path, max_chars)
    if ext == "pdf":
        return _read_pdf_text(path, max_chars)
    if ext == "docx":
        return _read_docx_text(path, max_chars)
    if ext in {"xlsx", "xlsm"}:
        return _read_xlsx_text(path, max_chars)
    return ""


def _read_existing_text_artifact(instance: InstanceConfig, row: dict[str, str], max_chars: int) -> str:
    text_path = row.get("text_path", "")
    if not text_path:
        return ""
    path = Path(text_path)
    if not path.is_absolute():
        path = instance.root("workspace") / path
    if not path.exists() or not path.is_file():
        return ""
    return _read_text_file(path, max_chars)


def _scan_roots(
    instance: InstanceConfig,
    include_generated: bool,
    *,
    scan_workspace_prefixes: bool | None = None,
) -> list[Path]:
    roots: list[Path] = []
    workspace = instance.root("workspace")
    privacy_settings = instance.settings().get("privacyops") or instance.settings().get("privacy") or {}
    if not isinstance(privacy_settings, dict):
        privacy_settings = {}
    try:
        roots.append(instance.root("raw"))
    except KeyError:
        pass
    if privacy_settings.get("scan_system"):
        try:
            roots.append(instance.root("system"))
        except KeyError:
            pass
    if include_generated:
        for root_name in ("outputs", "staging"):
            try:
                roots.append(instance.root(root_name))
            except KeyError:
                pass
    roots.extend(instance.root_list("restricted"))

    if scan_workspace_prefixes is None:
        scan_workspace_prefixes = bool(privacy_settings.get("scan_workspace_prefixes"))
    if scan_workspace_prefixes:
        for child in workspace.iterdir() if workspace.exists() else []:
            if child.is_dir() and child.name.startswith(DEFAULT_WORKSPACE_PREFIXES):
                roots.append(child)

    unique: list[Path] = []
    seen: set[str] = set()
    for root in roots:
        resolved = root.resolve()
        key = str(resolved).lower()
        if resolved.exists() and key not in seen:
            seen.add(key)
            unique.append(resolved)
    return unique


def _iter_files(roots: Iterable[Path]):
    for root in roots:
        for path in sorted(root.rglob("*")):
            if not path.is_file():
                continue
            if path.name.lower() in SKIP_FILE_NAMES:
                continue
            if any(marker in path.name.lower() for marker in SKIP_FILE_MARKERS):
                continue
            try:
                relative_parts = path.relative_to(root).parts
            except ValueError:
                relative_parts = path.parts
            if any(part.lower() in SKIP_DIR_NAMES for part in relative_parts):
                continue
            yield path


def _source_zone(workspace_root: Path, path: Path) -> str:
    try:
        rel = path.resolve().relative_to(workspace_root.resolve())
        return rel.parts[0] if rel.parts else "workspace"
    except ValueError:
        return path.parent.name or "external"


def _doc_id_for(digest: str) -> str:
    return f"DOC-{digest[:12].upper()}"


def _existing_indexes(rows: list[dict[str, str]]) -> tuple[dict[str, dict[str, str]], dict[str, dict[str, str]]]:
    by_path = {row.get("original_path", "").replace("\\", "/").lower(): row for row in rows if row.get("original_path")}
    by_digest = {row.get("sha256", ""): row for row in rows if row.get("sha256")}
    return by_path, by_digest


def _ensure_review_fields(fields: list[str]) -> list[str]:
    for field in REVIEW_DECISION_FIELDS:
        if field not in fields:
            fields.append(field)
    return fields


def _parts(value: str) -> set[str]:
    return {part.strip() for part in value.split(";") if part.strip()}


def _wide_diffusion_target(row: dict[str, str]) -> bool:
    target = normalize_college(
        row.get("derivative_max_college") or row.get("raw_max_college"),
        default="C2_Coproprietaires",
    )
    return ACCESS_RANK[target] <= ACCESS_RANK["C2_Coproprietaires"]


def _recommended_review_decision(row: dict[str, str]) -> tuple[str, str, str]:
    transformations = _parts(row.get("required_transformations", ""))
    publication_form = row.get("publication_form", "")
    priority, _exposure_risk, recommended_action = exposure_risk_for(row)

    if priority == "P0":
        return (
            "BLOQUE",
            "Retirer le document du chemin de diffusion large puis refaire une revue PrivacyOps.",
            "Document sensible detecte dans un chemin de diffusion large.",
        )
    if publication_form == "raw" and not transformations:
        return (
            "DIFFUSABLE_BRUT",
            "Valider que la source peut etre diffusee telle quelle et tracer la justification.",
            "",
        )
    if "aggregation" in transformations:
        return (
            "DIFFUSABLE_APRES_AGREGATION",
            "Produire une synthese agregee qui ne reprend pas les lignes sensibles.",
            "",
        )
    if "redaction" in transformations:
        return (
            "DIFFUSABLE_APRES_BIFFAGE",
            "Produire ou verifier une version biffee avant diffusion.",
            "",
        )
    if "metadata_only" in transformations or publication_form == "metadata_only":
        return (
            "BLOQUE",
            "Ne pas diffuser le contenu ; garder seulement une reference ou une meta-donnee.",
            "La politique limite ce document a une diffusion metadata_only.",
        )
    if row.get("review_required"):
        return (
            "A_ARBITRER",
            recommended_action or "Arbitrer le statut de diffusion avant toute sortie large.",
            "La politique automatique demande une revue humaine.",
        )
    return (
        "A_ARBITRER",
        recommended_action or "Verifier manuellement le statut de diffusion.",
        "Aucune decision de diffusion robuste n'a ete deduite.",
    )


def _has_human_review(previous: dict[str, str], digest: str) -> bool:
    if previous.get("sha256") != digest:
        return False
    status = previous.get("privacy_review_status", "")
    if status not in REVIEW_STATUSES:
        return False
    return bool(
        previous.get("privacy_reviewed_at")
        or previous.get("privacy_review_owner")
        or previous.get("privacy_review_justification")
    )


def _apply_review_defaults(row: dict[str, str], previous: dict[str, str], digest: str) -> None:
    recommendation, next_step, blocker = _recommended_review_decision(row)
    row["review_status_recommendation"] = recommendation
    row["review_next_step"] = next_step
    row["publication_blocker"] = blocker
    row["review_justification_required"] = "YES" if _wide_diffusion_target(row) else ""

    if _has_human_review(previous, digest):
        row["privacy_review_status"] = previous.get("privacy_review_status", "")
        row["privacy_review_justification"] = previous.get("privacy_review_justification", "")
        row["privacy_review_owner"] = previous.get("privacy_review_owner", "")
        row["privacy_reviewed_at"] = previous.get("privacy_reviewed_at", "")
        return

    if blocker:
        row["privacy_review_status"] = "BLOQUE"
    elif row.get("review_required"):
        row["privacy_review_status"] = "A_ARBITRER"
    else:
        row["privacy_review_status"] = recommendation
    row["privacy_review_justification"] = ""
    row["privacy_review_owner"] = ""
    row["privacy_reviewed_at"] = ""


def _validate_review_decision(row: dict[str, str], status: str, justification: str) -> None:
    transformations = _parts(row.get("required_transformations", ""))
    if row.get("publication_blocker") and status != "BLOQUE":
        raise ValueError(f"{row.get('doc_id', 'document')} is blocked by PrivacyOps: {row['publication_blocker']}")
    if status in DIFFUSION_REVIEW_STATUSES and _wide_diffusion_target(row) and not justification.strip():
        raise ValueError("A justification is required before any broad diffusion decision.")
    if status == "DIFFUSABLE_BRUT":
        blocking_transformations = transformations.intersection({"redaction", "aggregation", "metadata_only"})
        if row.get("publication_form") != "raw" or blocking_transformations:
            raise ValueError("Raw diffusion is not allowed while PrivacyOps requires redaction, aggregation or metadata_only.")


def _sync_screening_review_decision(instance: InstanceConfig, doc_id: str, updates: dict[str, str]) -> None:
    path = privacy_screening_path(instance)
    if not path.exists():
        return
    fields, rows = read_csv(path)
    fields = fields or list(SCREENING_REVIEW_FIELDS)
    for field in SCREENING_REVIEW_FIELDS:
        if field not in fields:
            fields.append(field)
    changed = False
    for row in rows:
        if row.get("doc_id") == doc_id:
            row.update(updates)
            changed = True
    if changed:
        write_csv(path, fields, rows)


def _base_document_row(
    instance: InstanceConfig,
    workspace_root: Path,
    path: Path,
    digest: str,
    previous: dict[str, str],
    seen_at: str,
) -> dict[str, str]:
    stat = path.stat()
    relative_path = relative_to(workspace_root, path)
    row = {field: "" for field in DEFAULT_DOCUMENT_FIELDS}
    row.update(previous)
    row.update(
        {
            "doc_id": previous.get("doc_id") or _doc_id_for(digest),
            "instance_id": instance.instance_id,
            "entity_id": instance.entity_id,
            "scope": previous.get("scope") or "screening",
            "sha256": digest,
            "original_path": relative_path,
            "file_name": path.name,
            "extension": path.suffix.lower().lstrip("."),
            "size_bytes": str(stat.st_size),
            "last_modified": datetime.fromtimestamp(stat.st_mtime, timezone.utc).replace(microsecond=0).isoformat(),
            "first_seen": previous.get("first_seen") or seen_at,
            "source_zone": previous.get("source_zone") or _source_zone(workspace_root, path),
            "source_kind": previous.get("source_kind") or "screened_existing",
            "status_ocr": previous.get("status_ocr") or "PENDING",
            "classification_status": previous.get("classification_status") or "PENDING",
        }
    )
    return row


def _screening_row(row: dict[str, str]) -> dict[str, str]:
    priority, exposure_risk, action = exposure_risk_for(row)
    values = {field: row.get(field, "") for field in SCREENING_REVIEW_FIELDS}
    values.update(
        {
            "exposure_risk": exposure_risk,
            "remediation_priority": priority,
            "recommended_action": action,
        }
    )
    return values
