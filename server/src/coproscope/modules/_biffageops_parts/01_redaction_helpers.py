from __future__ import annotations

import csv
import json
import re
from collections import Counter
from pathlib import Path

from ..core.common import (
    InstanceConfig,
    RunContext,
    now_iso,
    read_csv,
    relative_to,
    sha256_file,
    write_csv,
    write_text,
)
from ..core.privacy import ACCESS_RANK, PrivacySignal, normalize_college, redaction_signals
from ..core.sanitizer import extract_docx_package_text, sanitize_docx_package
from .privacyops import redactions_path


REDACTION_REGISTER_FIELDS = [
    "redaction_id",
    "created_at",
    "doc_id",
    "source_path",
    "source_sha256",
    "redacted_path",
    "redacted_sha256",
    "mode",
    "target_college",
    "status",
    "match_count",
    "categories",
    "map_path",
    "notes",
]

REDACTION_MAP_FIELDS = [
    "map_id",
    "created_at",
    "doc_id",
    "source_sha256",
    "category",
    "alias",
    "original_value",
    "reason",
    "storage_college",
    "retention_class",
]

REDACTION_QUEUE_FIELDS = [
    "doc_id",
    "file_name",
    "original_path",
    "raw_max_college",
    "derivative_max_college",
    "publication_form",
    "required_transformations",
    "review_required",
    "recommended_mode",
    "queue_status",
]

TEXT_EXTENSIONS = {"txt", "md", "csv", "tsv", "json", "yml", "yaml", "html", "htm"}
LOCAL_REDACTION_EXTENSIONS = TEXT_EXTENSIONS | {"pdf", "docx"}


def _redaction_extension(row: dict[str, str]) -> str:
    return Path(row.get("original_path") or row.get("file_name", "")).suffix.lower().lstrip(".")


def _supports_local_redaction(row: dict[str, str]) -> bool:
    return _redaction_extension(row) in LOCAL_REDACTION_EXTENSIONS


def _privacy_dir(instance: InstanceConfig) -> Path:
    path = instance.artifact("privacy_dir")
    path.mkdir(parents=True, exist_ok=True)
    return path


def _redacted_dir(instance: InstanceConfig) -> Path:
    path = instance.artifact("redacted_dir")
    path.mkdir(parents=True, exist_ok=True)
    return path


def _is_inside(path: Path, roots: list[Path]) -> bool:
    try:
        resolved = path.resolve()
    except OSError:
        resolved = path
    for root in roots:
        try:
            resolved.relative_to(root.resolve())
            return True
        except ValueError:
            continue
    return False


def _fallback_redaction_map_path(instance: InstanceConfig) -> Path:
    restricted_roots = instance.root_list("restricted")
    if restricted_roots:
        path = restricted_roots[0] / "biffageops" / "table_correspondance_biffage.csv"
    else:
        path = _privacy_dir(instance) / "C8_table_correspondance_biffage.csv"
    path.parent.mkdir(parents=True, exist_ok=True)
    return path


def _redaction_map_path(instance: InstanceConfig) -> Path:
    try:
        configured = instance.register("redaction_map")
    except KeyError:
        return _fallback_redaction_map_path(instance)

    restricted_roots = instance.root_list("restricted")
    configured_parts = {part.lower() for part in configured.parts}
    if (
        _is_inside(configured, restricted_roots)
        or configured_parts.intersection({"c8", "restreint", "restricted"})
        or configured.name.lower().startswith("c8_")
    ):
        configured.parent.mkdir(parents=True, exist_ok=True)
        return configured
    return _fallback_redaction_map_path(instance)


def _queue_path(instance: InstanceConfig) -> Path:
    return _privacy_dir(instance) / "file_biffage.csv"


def _source_path(instance: InstanceConfig, row: dict[str, str]) -> Path:
    value = row.get("original_path", "")
    path = Path(value)
    if path.is_absolute():
        return path
    return instance.root("workspace") / value


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _pdf_text(path: Path) -> str:
    try:
        import fitz  # type: ignore

        document = fitz.open(str(path))
        try:
            return "\n".join(document.load_page(index).get_text("text") for index in range(document.page_count))
        finally:
            document.close()
    except Exception:  # noqa: BLE001
        pass
    try:
        from pypdf import PdfReader  # type: ignore

        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception:  # noqa: BLE001
        return ""


def _docx_text(path: Path) -> str:
    parts: list[str] = []
    try:
        from docx import Document  # type: ignore

        document = Document(str(path))
        parts.extend(paragraph.text for paragraph in document.paragraphs)
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
    except Exception:  # noqa: BLE001
        pass
    try:
        parts.append(extract_docx_package_text(path))
    except Exception:  # noqa: BLE001
        pass
    return "\n".join(part for part in parts if part)


def _source_text(path: Path) -> str:
    ext = path.suffix.lower().lstrip(".")
    if ext in TEXT_EXTENSIONS:
        return _read_text(path)
    if ext == "pdf":
        return _pdf_text(path)
    if ext == "docx":
        return _docx_text(path)
    return ""


def _registered_text(instance: InstanceConfig, row: dict[str, str]) -> str:
    text_path = row.get("text_path", "")
    if not text_path:
        return ""
    path = Path(text_path)
    if not path.is_absolute():
        path = instance.root("workspace") / path
    if not path.exists():
        return ""
    return _read_text(path)


def _label_for(signal: PrivacySignal, mode: str, alias: str | None) -> str:
    if mode == "pseudonymisation_tracee" and alias:
        return alias
    return f"[BIFFE_{signal.category}]"


def _replace_values(text: str, replacements: dict[str, str]) -> str:
    redacted = text
    for original, replacement in sorted(replacements.items(), key=lambda item: len(item[0]), reverse=True):
        if not original:
            continue
        redacted = redacted.replace(original, replacement)
    return redacted


def _load_map(path: Path) -> tuple[list[str], list[dict[str, str]]]:
    fields, rows = read_csv(path)
    fields = fields or list(REDACTION_MAP_FIELDS)
    for field in REDACTION_MAP_FIELDS:
        if field not in fields:
            fields.append(field)
    return fields, rows


def _next_alias(category: str, rows: list[dict[str, str]]) -> str:
    prefix = {
        "PERSON_NAME": "PERSONNE",
        "EMAIL": "EMAIL",
        "PHONE": "TELEPHONE",
        "IBAN": "IBAN",
        "LOT": "LOT",
        "ACCOUNT_INDIVIDUAL": "COMPTE",
        "SECURITY_SECRET": "SECRET",
    }.get(category, category)
    existing = [
        row.get("alias", "")
        for row in rows
        if row.get("alias", "").startswith(prefix + "_")
    ]
    numbers: list[int] = []
    for alias in existing:
        match = re.search(r"_(\d+)$", alias)
        if match:
            numbers.append(int(match.group(1)))
    return f"{prefix}_{(max(numbers) if numbers else 0) + 1:04d}"


def _mapping_for(
    instance: InstanceConfig,
    row: dict[str, str],
    signals: list[PrivacySignal],
) -> tuple[dict[str, str], Path]:
    path = _redaction_map_path(instance)
    fields, rows = _load_map(path)
    now = now_iso()
    replacements: dict[str, str] = {}
    changed = False
    doc_id = row.get("doc_id", "")
    source_sha256 = row.get("sha256", "")
    for signal in signals:
        existing_alias = next(
            (
                item
                for item in rows
                if item.get("category") == signal.category and item.get("original_value") == signal.value
            ),
            None,
        )
        if existing_alias:
            alias = existing_alias.get("alias", "")
        else:
            alias = _next_alias(signal.category, rows)
        existing_doc_mapping = next(
            (
                item
                for item in rows
                if item.get("category") == signal.category
                and item.get("original_value") == signal.value
                and item.get("doc_id") == doc_id
                and item.get("source_sha256") == source_sha256
            ),
            None,
        )
        if existing_doc_mapping is None:
            rows.append(
                {
                    "map_id": f"MAP-{doc_id}-{signal.category}-{len(rows) + 1:04d}",
                    "created_at": now,
                    "doc_id": doc_id,
                    "source_sha256": source_sha256,
                    "category": signal.category,
                    "alias": alias,
                    "original_value": signal.value,
                    "reason": signal.reason,
                    "storage_college": "C8_Restreint_Critique",
                    "retention_class": "preuve_reidentification_limitee",
                }
            )
            changed = True
        replacements[signal.value] = alias
    if changed:
        write_csv(path, fields, rows)
    return replacements, path


def _text_output_path(instance: InstanceConfig, row: dict[str, str], source: Path) -> Path:
    suffix = source.suffix or ".txt"
    return _redacted_dir(instance) / f"{row.get('doc_id', 'DOC')}.redacted{suffix}"


def _write_text_redaction(
    instance: InstanceConfig,
    row: dict[str, str],
    source: Path,
    signals: list[PrivacySignal],
    replacements: dict[str, str],
    mode: str,
) -> Path:
    text = _read_text(source)
    if mode == "redaction_irreversible":
        replacements = {signal.value: _label_for(signal, mode, None) for signal in signals}
    out = _text_output_path(instance, row, source)
    write_text(out, _replace_values(text, replacements))
    return out


def _write_text_derivative_redaction(
    instance: InstanceConfig,
    row: dict[str, str],
    text: str,
    signals: list[PrivacySignal],
    replacements: dict[str, str],
    mode: str,
) -> Path:
    if mode == "redaction_irreversible":
        replacements = {signal.value: _label_for(signal, mode, None) for signal in signals}
    out = _redacted_dir(instance) / f"{row.get('doc_id', 'DOC')}.redacted.txt"
    write_text(out, _replace_values(text, replacements))
    return out


def _write_pdf_redaction(
    instance: InstanceConfig,
    row: dict[str, str],
    source: Path,
    signals: list[PrivacySignal],
    replacements: dict[str, str],
    mode: str,
) -> Path:
    try:
        import fitz  # type: ignore
    except ImportError as exc:
        raise RuntimeError("PyMuPDF/fitz unavailable for PDF redaction.") from exc

    out = _redacted_dir(instance) / f"{row.get('doc_id', 'DOC')}.redacted.pdf"
    document = fitz.open(str(source))
    try:
        for page in document:
            for signal in signals:
                label = _label_for(signal, mode, replacements.get(signal.value))
                for rect in page.search_for(signal.value):
                    page.add_redact_annot(rect, text=label, fill=(0, 0, 0), text_color=(1, 1, 1))
            page.apply_redactions()
        document.set_metadata({})
        document.save(str(out), garbage=4, deflate=True, clean=True)
    finally:
        document.close()
    return out


def _replace_paragraph_text(paragraph, replacements: dict[str, str]) -> None:  # type: ignore[no-untyped-def]
    original = paragraph.text
    redacted = _replace_values(original, replacements)
    if redacted == original:
        return
    if paragraph.runs:
        paragraph.runs[0].text = redacted
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(redacted)


def _write_docx_redaction(
    instance: InstanceConfig,
    row: dict[str, str],
    source: Path,
    signals: list[PrivacySignal],
    replacements: dict[str, str],
    mode: str,
) -> Path:
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise RuntimeError("python-docx unavailable for DOCX redaction.") from exc

    if mode == "redaction_irreversible":
        replacements = {signal.value: _label_for(signal, mode, None) for signal in signals}
    document = Document(str(source))
    for paragraph in document.paragraphs:
        _replace_paragraph_text(paragraph, replacements)
    for table in document.tables:
        for table_row in table.rows:
            for cell in table_row.cells:
                for paragraph in cell.paragraphs:
                    _replace_paragraph_text(paragraph, replacements)
    document.core_properties.author = ""
    document.core_properties.comments = ""
    document.core_properties.keywords = ""
    document.core_properties.subject = ""
    document.core_properties.title = ""
    out = _redacted_dir(instance) / f"{row.get('doc_id', 'DOC')}.redacted.docx"
    document.save(str(out))
    sanitize_docx_package(out, replacements=replacements)
    return out


def _write_redacted(
    instance: InstanceConfig,
    row: dict[str, str],
    source: Path,
    signals: list[PrivacySignal],
    replacements: dict[str, str],
    mode: str,
    *,
    fallback_text: str | None = None,
) -> Path:
    ext = source.suffix.lower().lstrip(".")
    if ext in TEXT_EXTENSIONS:
        return _write_text_redaction(instance, row, source, signals, replacements, mode)
    if ext == "pdf":
        if fallback_text is not None and not _pdf_text(source).strip():
            return _write_text_derivative_redaction(instance, row, fallback_text, signals, replacements, mode)
        try:
            return _write_pdf_redaction(instance, row, source, signals, replacements, mode)
        except RuntimeError:
            if fallback_text is not None:
                return _write_text_derivative_redaction(instance, row, fallback_text, signals, replacements, mode)
            raise
    if ext == "docx":
        return _write_docx_redaction(instance, row, source, signals, replacements, mode)
    raise RuntimeError(f"Unsupported redaction format: {ext or 'no_extension'}")


def _register_redaction(instance: InstanceConfig, record: dict[str, str]) -> None:
    path = redactions_path(instance)
    fields, rows = read_csv(path)
    fields = fields or list(REDACTION_REGISTER_FIELDS)
    for field in REDACTION_REGISTER_FIELDS:
        if field not in fields:
            fields.append(field)
    rows = [row for row in rows if row.get("redaction_id") != record.get("redaction_id")]
    rows.append(record)
    write_csv(path, fields, rows)


def _update_document_row(instance: InstanceConfig, doc_id: str, updates: dict[str, str]) -> None:
    documents_path = instance.register("documents")
    fields, rows = read_csv(documents_path)
    for field in updates:
        if field not in fields:
            fields.append(field)
    for row in rows:
        if row.get("doc_id") == doc_id:
            row.update(updates)
    write_csv(documents_path, fields, rows)
