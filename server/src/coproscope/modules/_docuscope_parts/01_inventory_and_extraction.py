from __future__ import annotations

import csv
import re
import shutil
from collections import defaultdict
from datetime import date, datetime, timezone
from html.parser import HTMLParser
from pathlib import Path

from ..core.common import (
    DEFAULT_DOCUMENT_FIELDS,
    FINDING_FIELDS,
    RunContext,
    append_note,
    detect_date,
    load_template_csv,
    load_structured_file,
    now_iso,
    read_csv,
    relative_to,
    safe_name,
    sha256_file,
    write_csv,
    write_text,
)
from ..core.privacy import apply_access_policy, ensure_policy_fields
from .coprolink import request_metrics


TEXT_EXTENSIONS = {"txt", "md", "csv", "tsv", "json"}
HTML_EXTENSIONS = {"html", "htm"}
SKIP_DIRS = {
    ".git",
    ".venv",
    "__pycache__",
    ".secrets",
    ".pytest_cache",
    ".mypy_cache",
    ".ruff_cache",
    ".cache",
    ".coproscope",
    "node_modules",
    "coproscope",
    "server",
    "docs",
    "00 - lire avant de partager",
    "08 - versions partageables",
    "09 - rapports et syntheses",
    "_archives",
    "_instance_docs",
    "_transition_reports",
    "_worktrees",
    "900_systeme_audit",
    "990_archives",
}
SKIP_FILENAMES = {"desktop.ini", "thumbs.db", ".gitignore", "instance.yml", "passation_coproscope_local.md"}
PRESERVE_FIELDS = {
    "lot",
    "document_type",
    "suspected_date",
    "emitter",
    "status_ocr",
    "classification_status",
    "text_path",
    "page_count",
    "text_char_count",
    "extraction_level",
    "text_quality",
    "ocr_engine",
    "docling_path",
    "layout_path",
    "ai_review_status",
    "sensitivity",
    "raw_max_college",
    "derivative_max_college",
    "publication_form",
    "restriction_reasons",
    "personal_data_level",
    "ip_status",
    "ai_processing_ceiling",
    "review_required",
    "policy_confidence",
    "required_transformations",
    "screening_signals",
    "privacy_review_status",
    "redaction_status",
    "redaction_mode",
    "redacted_path",
    "redacted_sha256",
    "redaction_map_id",
    "notes",
}
COMPLETENESS_FIELDS = [
    "proof_id",
    "lot",
    "expected_label",
    "document_type",
    "status",
    "criticality",
    "freshness_months",
    "matched_doc_ids",
    "evidence_paths",
    "newest_date",
    "reason",
    "action",
]
DOCUMENT_REQUEST_FIELDS = [
    "request_id",
    "source_ref",
    "priority",
    "status",
    "subject",
    "expected_piece",
    "reason",
    "related_doc_ids",
    "evidence_paths",
    "suggested_diligence",
]


def _load_existing_by_digest(registry_path: Path) -> dict[str, dict[str, str]]:
    _, rows = read_csv(registry_path)
    return {row.get("sha256", ""): row for row in rows if row.get("sha256")}


def _doc_id_for(digest: str) -> str:
    return f"DOC-{digest[:12].upper()}"


def _is_relative_to(path: Path, parent: Path) -> bool:
    try:
        path.resolve().relative_to(parent.resolve())
        return True
    except ValueError:
        return False


def _raw_roots(instance) -> list[Path]:
    roots = instance.root_list("raw")
    return roots or [instance.root("raw")]


def _dedupe_scan_roots(roots: list[Path]) -> list[Path]:
    result: list[Path] = []
    for root in [item.resolve() for item in roots]:
        if any(_is_relative_to(root, existing) for existing in result):
            continue
        result = [existing for existing in result if not _is_relative_to(existing, root)]
        result.append(root)
    return result


def _configured_inventory_roots(instance, raw_roots: list[Path]) -> list[Path]:
    roots = list(raw_roots)
    physical_deposit = instance.physical_deposit_path()
    if physical_deposit is not None:
        roots.append(physical_deposit)
    if instance.user_moves_allowed():
        roots.append(instance.root("workspace"))
    return _dedupe_scan_roots(roots)


def _protected_roots(instance) -> list[Path]:
    roots: list[Path] = []
    for name in ["system", "outputs", "staging", "logs"]:
        try:
            roots.append(instance.root(name))
        except KeyError:
            pass
    technical_root = instance.technical_root()
    if technical_root is not None:
        roots.append(technical_root)
    for name in ["documents", "duplicates", "manifest", "requests", "ag", "findings", "kpi"]:
        try:
            roots.append(instance.register(name).parent)
        except KeyError:
            pass
    return _dedupe_scan_roots(roots)


def _is_protected(path: Path, scan_root: Path, protected_roots: list[Path]) -> bool:
    return any(
        _is_relative_to(path, protected) and not _is_relative_to(scan_root, protected)
        for protected in protected_roots
    )


def _source_labels(path: Path, instance, raw_roots: list[Path]) -> tuple[str, str]:
    physical_deposit = instance.physical_deposit_path()
    if physical_deposit is not None and _is_relative_to(path, physical_deposit):
        return "DEPOT_PHYSIQUE", "physical_deposit"
    if any(_is_relative_to(path, raw_root) for raw_root in raw_roots):
        return "RAW", "raw"
    if instance.user_moves_allowed() and _is_relative_to(path, instance.root("workspace")):
        return "COFFRE_VISIBLE", "visible_vault"
    return "RAW", "raw"


def _iter_files(root: Path, protected_roots: list[Path]):
    for path in sorted(root.rglob("*")):
        if not path.is_file():
            continue
        if _is_protected(path, root, protected_roots):
            continue
        try:
            relative_parts = path.relative_to(root).parts
        except ValueError:
            relative_parts = path.parts
        if any(part.lower() in SKIP_DIRS for part in relative_parts):
            continue
        if path.name.lower() in SKIP_FILENAMES:
            continue
        yield path


def _load_sensitivity_rules(instance) -> dict:
    return load_structured_file(instance.sensitivity_rules_path())


def _classify_sensitivity(relative_path: str, instance) -> tuple[str, str]:
    rules = _load_sensitivity_rules(instance)
    lower_path = relative_path.lower()
    for rule in rules.get("rules", []):
        keywords = [keyword.lower() for keyword in rule.get("path_keywords", [])]
        if any(keyword in lower_path for keyword in keywords):
            return str(rule.get("sensitivity", "internal")), str(rule.get("scope", "internal"))
    return str(rules.get("default_sensitivity", "internal")), "internal"


def inventory(instance, run: RunContext) -> Path:
    raw_roots = _raw_roots(instance)
    scan_roots = _configured_inventory_roots(instance, raw_roots)
    protected_roots = _protected_roots(instance)
    workspace_root = instance.root("workspace")
    registry_path = instance.register("documents")
    duplicates_path = instance.register("duplicates")
    manifest_path = instance.register("manifest")
    existing = _load_existing_by_digest(registry_path)
    seen_at = now_iso()
    rows: list[dict[str, str]] = []
    by_digest: dict[str, list[dict[str, str]]] = defaultdict(list)

    for scan_root in scan_roots:
        if not scan_root.exists():
            run.log_error(f"inventory source missing: {scan_root}")
            continue
        for path in _iter_files(scan_root, protected_roots):
            digest = sha256_file(path)
            previous = existing.get(digest, {})
            stat = path.stat()
            relative_path = relative_to(workspace_root, path)
            sensitivity, scope = _classify_sensitivity(relative_path, instance)
            source_zone, source_kind = _source_labels(path, instance, raw_roots)
            row = {field: "" for field in DEFAULT_DOCUMENT_FIELDS}
            row.update(
                {
                    "doc_id": previous.get("doc_id") or _doc_id_for(digest),
                    "instance_id": instance.instance_id,
                    "entity_id": instance.entity_id,
                    "scope": previous.get("scope") or scope,
                    "sha256": digest,
                    "original_path": relative_path,
                    "file_name": path.name,
                    "extension": path.suffix.lower().lstrip("."),
                    "size_bytes": str(stat.st_size),
                    "last_modified": datetime.fromtimestamp(stat.st_mtime, timezone.utc).replace(microsecond=0).isoformat(),
                    "first_seen": previous.get("first_seen") or seen_at,
                    "source_zone": source_zone,
                    "source_kind": source_kind,
                    "status_ocr": previous.get("status_ocr") or "PENDING",
                    "classification_status": previous.get("classification_status") or "PENDING",
                    "sensitivity": previous.get("sensitivity") or sensitivity,
                }
            )
            for field in PRESERVE_FIELDS:
                if previous.get(field):
                    row[field] = previous[field]
            apply_access_policy(row, instance=instance)
            rows.append(row)
            by_digest[digest].append(row)

    rows.sort(key=lambda item: (item["file_name"].lower(), item["doc_id"]))
    write_csv(registry_path, ensure_policy_fields(list(DEFAULT_DOCUMENT_FIELDS)), rows)
    run.log_action("write", registry_path, f"inventory rows={len(rows)}")

    duplicates_rows: list[dict[str, str]] = []
    for digest, digest_rows in sorted(by_digest.items()):
        if len(digest_rows) < 2:
            continue
        duplicates_rows.append(
            {
                "sha256": digest,
                "count": str(len(digest_rows)),
                "doc_ids": "; ".join(row["doc_id"] for row in digest_rows),
                "paths": "; ".join(row["original_path"] for row in digest_rows),
            }
        )
    write_csv(duplicates_path, ["sha256", "count", "doc_ids", "paths"], duplicates_rows)
    run.log_action("write", duplicates_path, f"duplicate groups={len(duplicates_rows)}")

    manifest_rows = [
        {
            "doc_id": row["doc_id"],
            "sha256": row["sha256"],
            "original_path": row["original_path"],
            "file_name": row["file_name"],
            "size_bytes": row["size_bytes"],
            "source_kind": row["source_kind"],
        }
        for row in rows
    ]
    write_csv(manifest_path, ["doc_id", "sha256", "original_path", "file_name", "size_bytes", "source_kind"], manifest_rows)
    run.log_action("write", manifest_path, f"manifest rows={len(manifest_rows)}")
    return registry_path


def _read_rows(path: Path) -> tuple[list[str], list[dict[str, str]]]:
    fields, rows = read_csv(path)
    if not rows and not fields:
        raise SystemExit(f"Missing registry: {path}")
    return fields, rows


def _write_rows(path: Path, fields: list[str], rows: list[dict[str, str]]) -> None:
    for field in [
        "status_ocr",
        "text_path",
        "page_count",
        "text_char_count",
        "extraction_level",
        "text_quality",
        "ocr_engine",
        "docling_path",
        "layout_path",
        "ai_review_status",
        "notes",
    ]:
        if field not in fields:
            fields.append(field)
    ensure_policy_fields(fields)
    write_csv(path, fields, rows)


def _extract_pdf_with_pymupdf(path: Path) -> tuple[list[str], str]:
    try:
        import fitz  # type: ignore
    except ImportError as exc:
        raise RuntimeError("Module PyMuPDF/fitz indisponible.") from exc

    document = fitz.open(str(path))
    try:
        pages = [document.load_page(index).get_text("text").strip() for index in range(document.page_count)]
        return pages, str(document.page_count)
    finally:
        document.close()


def _extract_pdf_with_pypdf(path: Path) -> tuple[list[str], str]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("Module pypdf indisponible.") from exc

    reader = PdfReader(str(path))
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:  # noqa: BLE001
            text = f"[ERREUR_EXTRACTION_PAGE_{index}: {exc}]"
        pages.append(text.strip())
    return pages, str(len(reader.pages))


def _extract_pdf(path: Path) -> tuple[list[str], str, str]:
    try:
        pages, page_count = _extract_pdf_with_pymupdf(path)
        return pages, page_count, "L1_NATIVE_TEXT_PYMUPDF"
    except Exception:  # noqa: BLE001
        pages, page_count = _extract_pdf_with_pypdf(path)
        return pages, page_count, "L1_NATIVE_TEXT_PYPDF"


def _extract_plain_text(path: Path) -> tuple[list[str], str]:
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = path.read_text(encoding="latin-1")
    return [text], "1"


class _VisibleTextParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []
        self.skip_depth = 0

    def handle_starttag(self, tag: str, attrs) -> None:  # noqa: ANN001
        if tag.lower() in {"script", "style", "noscript"}:
            self.skip_depth += 1
        if tag.lower() in {"p", "br", "div", "li", "tr", "h1", "h2", "h3", "h4"}:
            self.parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag.lower() in {"script", "style", "noscript"} and self.skip_depth:
            self.skip_depth -= 1
        if tag.lower() in {"p", "div", "li", "tr", "h1", "h2", "h3", "h4"}:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        if not self.skip_depth:
            self.parts.append(data)

    def text(self) -> str:
        return re.sub(r"\n{3,}", "\n\n", re.sub(r"[ \t]+", " ", "".join(self.parts))).strip()


def _extract_html(path: Path) -> tuple[list[str], str]:
    try:
        raw = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        raw = path.read_text(encoding="latin-1")
    parser = _VisibleTextParser()
    parser.feed(raw)
    return [parser.text()], "1"


def _extract_docx(path: Path) -> tuple[list[str], str]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("Module python-docx indisponible.") from exc

    document = Document(str(path))
    lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                lines.append(" | ".join(values))
    return ["\n".join(lines)], "1"


def _extract_xlsx(path: Path) -> tuple[list[str], str]:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RuntimeError("Module openpyxl indisponible.") from exc

    workbook = load_workbook(str(path), read_only=True, data_only=True)
    try:
        pages: list[str] = []
        for worksheet in workbook.worksheets:
            lines = [f"===== FEUILLE {worksheet.title} ====="]
            for row in worksheet.iter_rows(values_only=True):
                values = [str(value).strip() if value is not None else "" for value in row]
                if any(values):
                    lines.append(" | ".join(values).rstrip())
            pages.append("\n".join(lines))
    finally:
        workbook.close()
    return pages, str(len(pages))


def _write_page_text(text_dir: Path, doc_id: str, pages: list[str], suffix: str = "native") -> Path:
    text_dir.mkdir(parents=True, exist_ok=True)
    out = text_dir / f"{doc_id}.{suffix}.txt"
    lines: list[str] = []
    for index, text in enumerate(pages, start=1):
        lines.append(f"===== PAGE {index} =====")
        lines.append(text)
        lines.append("")
    out.write_text("\n".join(lines), encoding="utf-8")
    return out


def _process_row(instance, row: dict[str, str]) -> dict[str, str]:
    workspace_root = instance.root("workspace")
    text_dir = instance.artifact("text_dir")
    path = workspace_root / row["original_path"]
    ext = row.get("extension", "").lower()
    doc_id = row["doc_id"]
    row.setdefault("notes", "")

    if not path.exists():
        row["status_ocr"] = "SOURCE_MISSING"
        row["notes"] = append_note(row.get("notes", ""), "Fichier source introuvable.")
        return row

    try:
        if ext == "pdf":
            pages, page_count, extraction_level = _extract_pdf(path)
        elif ext in TEXT_EXTENSIONS:
            pages, page_count = _extract_plain_text(path)
            extraction_level = "L1_NATIVE_TEXT"
        elif ext in HTML_EXTENSIONS:
            pages, page_count = _extract_html(path)
            extraction_level = "L1_NATIVE_HTML"
        elif ext == "docx":
            pages, page_count = _extract_docx(path)
            extraction_level = "L1_NATIVE_DOCX"
        elif ext == "xlsx":
            pages, page_count = _extract_xlsx(path)
            extraction_level = "L1_NATIVE_XLSX"
        else:
            row["status_ocr"] = "OCR_REQUIRED" if ext in {"png", "jpg", "jpeg", "tif", "tiff", "bmp"} else "UNSUPPORTED"
            row["extraction_level"] = "L2_LOCAL_OCR_REQUIRED" if row["status_ocr"] == "OCR_REQUIRED" else ""
            row["text_quality"] = "missing" if row["status_ocr"] == "OCR_REQUIRED" else "unsupported"
            return row
    except Exception as exc:  # noqa: BLE001
        row["status_ocr"] = "EXTRACTION_ERROR"
        row["notes"] = append_note(row.get("notes", ""), f"Extraction impossible: {exc}")
        return row

    total_chars = sum(len(page.strip()) for page in pages)
    out = _write_page_text(text_dir, doc_id, pages)
    row["text_path"] = relative_to(workspace_root, out)
    row["page_count"] = page_count
    row["text_char_count"] = str(total_chars)
    row["extraction_level"] = extraction_level
    row["text_quality"] = "strong" if total_chars >= 80 else "weak"
    if total_chars >= 80 or ext != "pdf":
        row["status_ocr"] = "TEXT_EXTRACTED"
        if total_chars < 80:
            row["notes"] = append_note(row.get("notes", ""), "Texte court extrait; OCR non applicable a ce format.")
    else:
        row["status_ocr"] = "OCR_REQUIRED"
        row["notes"] = append_note(row.get("notes", ""), "Texte natif faible ou absent: OCR scan requis.")
    return row


def extract_text(instance, run: RunContext, doc_id: str | None = None, docai_mode: str | None = None) -> Path:
    registry_path = instance.register("documents")
    fields, rows = _read_rows(registry_path)
    processed = 0
    for row in rows:
        if doc_id and row.get("doc_id") != doc_id:
            continue
        if row.get("status_ocr") == "OCR_DONE":
            continue
        row.update(_process_row(instance, row))
        text = _text_sample(instance, row) if row.get("text_path") else ""
        apply_access_policy(row, text=text, instance=instance)
        processed += 1
    _write_rows(registry_path, fields, rows)
    run.log_action("write", registry_path, f"text extraction processed={processed}")
    if docai_mode and docai_mode != "off":
        from . import docai

        docai.process_after_native_extract(instance, run, doc_id=doc_id, mode=docai_mode)
    return registry_path


def _load_taxonomy(instance) -> dict:
    return load_structured_file(instance.classification_rules_path())


def _text_sample(instance, row: dict[str, str]) -> str:
    workspace_root = instance.root("workspace")
    parts = [row.get("file_name", ""), row.get("original_path", "")]
    text_path = row.get("text_path")
    if text_path:
        path = workspace_root / text_path
        if path.exists():
            parts.append(path.read_text(encoding="utf-8", errors="ignore")[:6000])
    docling_path = row.get("docling_path")
    if docling_path:
        path = workspace_root / docling_path
        if path.exists():
            parts.append(path.read_text(encoding="utf-8", errors="ignore")[:6000])
    return "\n".join(parts).lower()
