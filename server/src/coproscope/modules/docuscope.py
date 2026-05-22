from __future__ import annotations

import csv
import re
import shutil
from collections import defaultdict
from datetime import date, datetime, timezone
from html.parser import HTMLParser
from pathlib import Path

from ..core.common import (
    DEFAULT_DOCUMENT_FIELDS,
    FINDING_FIELDS,
    RunContext,
    append_note,
    detect_date,
    load_template_csv,
    load_structured_file,
    now_iso,
    read_csv,
    relative_to,
    safe_name,
    sha256_file,
    write_csv,
    write_text,
)
from ..core.privacy import apply_access_policy, ensure_policy_fields
from .coprolink import request_metrics


TEXT_EXTENSIONS = {"txt", "md", "csv", "tsv", "json"}
HTML_EXTENSIONS = {"html", "htm"}
SKIP_DIRS = {
    ".git",
    ".venv",
    "__pycache__",
    ".secrets",
    ".pytest_cache",
    ".mypy_cache",
    ".ruff_cache",
    ".cache",
    "node_modules",
    "coproscope",
    "server",
    "docs",
    "_archives",
    "_instance_docs",
    "_transition_reports",
    "_worktrees",
    "900_systeme_audit",
    "990_archives",
}
SKIP_FILENAMES = {"desktop.ini", "thumbs.db", "passation_coproscope_local.md"}
PRESERVE_FIELDS = {
    "lot",
    "document_type",
    "suspected_date",
    "emitter",
    "status_ocr",
    "classification_status",
    "text_path",
    "page_count",
    "text_char_count",
    "extraction_level",
    "text_quality",
    "ocr_engine",
    "docling_path",
    "layout_path",
    "ai_review_status",
    "sensitivity",
    "raw_max_college",
    "derivative_max_college",
    "publication_form",
    "restriction_reasons",
    "personal_data_level",
    "ip_status",
    "ai_processing_ceiling",
    "review_required",
    "policy_confidence",
    "required_transformations",
    "screening_signals",
    "privacy_review_status",
    "redaction_status",
    "redaction_mode",
    "redacted_path",
    "redacted_sha256",
    "redaction_map_id",
    "notes",
}
COMPLETENESS_FIELDS = [
    "proof_id",
    "lot",
    "expected_label",
    "document_type",
    "status",
    "criticality",
    "freshness_months",
    "matched_doc_ids",
    "evidence_paths",
    "newest_date",
    "reason",
    "action",
]
DOCUMENT_REQUEST_FIELDS = [
    "request_id",
    "source_ref",
    "priority",
    "status",
    "subject",
    "expected_piece",
    "reason",
    "related_doc_ids",
    "evidence_paths",
    "suggested_diligence",
]


def _load_existing_by_digest(registry_path: Path) -> dict[str, dict[str, str]]:
    _, rows = read_csv(registry_path)
    return {row.get("sha256", ""): row for row in rows if row.get("sha256")}


def _doc_id_for(digest: str) -> str:
    return f"DOC-{digest[:12].upper()}"


def _iter_files(root: Path):
    for path in sorted(root.rglob("*")):
        if not path.is_file():
            continue
        try:
            relative_parts = path.relative_to(root).parts
        except ValueError:
            relative_parts = path.parts
        if any(part.lower() in SKIP_DIRS for part in relative_parts):
            continue
        if path.name.lower() in SKIP_FILENAMES:
            continue
        yield path


def _load_sensitivity_rules(instance) -> dict:
    return load_structured_file(instance.sensitivity_rules_path())


def _classify_sensitivity(relative_path: str, instance) -> tuple[str, str]:
    rules = _load_sensitivity_rules(instance)
    lower_path = relative_path.lower()
    for rule in rules.get("rules", []):
        keywords = [keyword.lower() for keyword in rule.get("path_keywords", [])]
        if any(keyword in lower_path for keyword in keywords):
            return str(rule.get("sensitivity", "internal")), str(rule.get("scope", "internal"))
    return str(rules.get("default_sensitivity", "internal")), "internal"


def inventory(instance, run: RunContext) -> Path:
    raw_roots = instance.root_list("raw")
    if not raw_roots:
        raw_roots = [instance.root("raw")]
    workspace_root = instance.root("workspace")
    registry_path = instance.register("documents")
    duplicates_path = instance.register("duplicates")
    manifest_path = instance.register("manifest")
    existing = _load_existing_by_digest(registry_path)
    seen_at = now_iso()
    rows: list[dict[str, str]] = []
    by_digest: dict[str, list[dict[str, str]]] = defaultdict(list)

    for raw_root in raw_roots:
        if not raw_root.exists():
            run.log_error(f"inventory source missing: {raw_root}")
            continue
        for path in _iter_files(raw_root):
            digest = sha256_file(path)
            previous = existing.get(digest, {})
            stat = path.stat()
            relative_path = relative_to(workspace_root, path)
            sensitivity, scope = _classify_sensitivity(relative_path, instance)
            row = {field: "" for field in DEFAULT_DOCUMENT_FIELDS}
            row.update(
                {
                    "doc_id": previous.get("doc_id") or _doc_id_for(digest),
                    "instance_id": instance.instance_id,
                    "entity_id": instance.entity_id,
                    "scope": previous.get("scope") or scope,
                    "sha256": digest,
                    "original_path": relative_path,
                    "file_name": path.name,
                    "extension": path.suffix.lower().lstrip("."),
                    "size_bytes": str(stat.st_size),
                    "last_modified": datetime.fromtimestamp(stat.st_mtime, timezone.utc).replace(microsecond=0).isoformat(),
                    "first_seen": previous.get("first_seen") or seen_at,
                    "source_zone": "RAW",
                    "source_kind": "raw",
                    "status_ocr": previous.get("status_ocr") or "PENDING",
                    "classification_status": previous.get("classification_status") or "PENDING",
                    "sensitivity": previous.get("sensitivity") or sensitivity,
                }
            )
            for field in PRESERVE_FIELDS:
                if previous.get(field):
                    row[field] = previous[field]
            apply_access_policy(row, instance=instance)
            rows.append(row)
            by_digest[digest].append(row)

    rows.sort(key=lambda item: (item["file_name"].lower(), item["doc_id"]))
    write_csv(registry_path, ensure_policy_fields(list(DEFAULT_DOCUMENT_FIELDS)), rows)
    run.log_action("write", registry_path, f"inventory rows={len(rows)}")

    duplicates_rows: list[dict[str, str]] = []
    for digest, digest_rows in sorted(by_digest.items()):
        if len(digest_rows) < 2:
            continue
        duplicates_rows.append(
            {
                "sha256": digest,
                "count": str(len(digest_rows)),
                "doc_ids": "; ".join(row["doc_id"] for row in digest_rows),
                "paths": "; ".join(row["original_path"] for row in digest_rows),
            }
        )
    write_csv(duplicates_path, ["sha256", "count", "doc_ids", "paths"], duplicates_rows)
    run.log_action("write", duplicates_path, f"duplicate groups={len(duplicates_rows)}")

    manifest_rows = [
        {
            "doc_id": row["doc_id"],
            "sha256": row["sha256"],
            "original_path": row["original_path"],
            "file_name": row["file_name"],
            "size_bytes": row["size_bytes"],
            "source_kind": row["source_kind"],
        }
        for row in rows
    ]
    write_csv(manifest_path, ["doc_id", "sha256", "original_path", "file_name", "size_bytes", "source_kind"], manifest_rows)
    run.log_action("write", manifest_path, f"manifest rows={len(manifest_rows)}")
    return registry_path


def _read_rows(path: Path) -> tuple[list[str], list[dict[str, str]]]:
    fields, rows = read_csv(path)
    if not rows and not fields:
        raise SystemExit(f"Missing registry: {path}")
    return fields, rows


def _write_rows(path: Path, fields: list[str], rows: list[dict[str, str]]) -> None:
    for field in [
        "status_ocr",
        "text_path",
        "page_count",
        "text_char_count",
        "extraction_level",
        "text_quality",
        "ocr_engine",
        "docling_path",
        "layout_path",
        "ai_review_status",
        "notes",
    ]:
        if field not in fields:
            fields.append(field)
    ensure_policy_fields(fields)
    write_csv(path, fields, rows)


def _extract_pdf_with_pymupdf(path: Path) -> tuple[list[str], str]:
    try:
        import fitz  # type: ignore
    except ImportError as exc:
        raise RuntimeError("Module PyMuPDF/fitz indisponible.") from exc

    document = fitz.open(str(path))
    try:
        pages = [document.load_page(index).get_text("text").strip() for index in range(document.page_count)]
        return pages, str(document.page_count)
    finally:
        document.close()


def _extract_pdf_with_pypdf(path: Path) -> tuple[list[str], str]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("Module pypdf indisponible.") from exc

    reader = PdfReader(str(path))
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:  # noqa: BLE001
            text = f"[ERREUR_EXTRACTION_PAGE_{index}: {exc}]"
        pages.append(text.strip())
    return pages, str(len(reader.pages))


def _extract_pdf(path: Path) -> tuple[list[str], str, str]:
    try:
        pages, page_count = _extract_pdf_with_pymupdf(path)
        return pages, page_count, "L1_NATIVE_TEXT_PYMUPDF"
    except Exception:  # noqa: BLE001
        pages, page_count = _extract_pdf_with_pypdf(path)
        return pages, page_count, "L1_NATIVE_TEXT_PYPDF"


def _extract_plain_text(path: Path) -> tuple[list[str], str]:
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = path.read_text(encoding="latin-1")
    return [text], "1"


class _VisibleTextParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []
        self.skip_depth = 0

    def handle_starttag(self, tag: str, attrs) -> None:  # noqa: ANN001
        if tag.lower() in {"script", "style", "noscript"}:
            self.skip_depth += 1
        if tag.lower() in {"p", "br", "div", "li", "tr", "h1", "h2", "h3", "h4"}:
            self.parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag.lower() in {"script", "style", "noscript"} and self.skip_depth:
            self.skip_depth -= 1
        if tag.lower() in {"p", "div", "li", "tr", "h1", "h2", "h3", "h4"}:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        if not self.skip_depth:
            self.parts.append(data)

    def text(self) -> str:
        return re.sub(r"\n{3,}", "\n\n", re.sub(r"[ \t]+", " ", "".join(self.parts))).strip()


def _extract_html(path: Path) -> tuple[list[str], str]:
    try:
        raw = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        raw = path.read_text(encoding="latin-1")
    parser = _VisibleTextParser()
    parser.feed(raw)
    return [parser.text()], "1"


def _extract_docx(path: Path) -> tuple[list[str], str]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("Module python-docx indisponible.") from exc

    document = Document(str(path))
    lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                lines.append(" | ".join(values))
    return ["\n".join(lines)], "1"


def _extract_xlsx(path: Path) -> tuple[list[str], str]:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RuntimeError("Module openpyxl indisponible.") from exc

    workbook = load_workbook(str(path), read_only=True, data_only=True)
    try:
        pages: list[str] = []
        for worksheet in workbook.worksheets:
            lines = [f"===== FEUILLE {worksheet.title} ====="]
            for row in worksheet.iter_rows(values_only=True):
                values = [str(value).strip() if value is not None else "" for value in row]
                if any(values):
                    lines.append(" | ".join(values).rstrip())
            pages.append("\n".join(lines))
    finally:
        workbook.close()
    return pages, str(len(pages))


def _write_page_text(text_dir: Path, doc_id: str, pages: list[str], suffix: str = "native") -> Path:
    text_dir.mkdir(parents=True, exist_ok=True)
    out = text_dir / f"{doc_id}.{suffix}.txt"
    lines: list[str] = []
    for index, text in enumerate(pages, start=1):
        lines.append(f"===== PAGE {index} =====")
        lines.append(text)
        lines.append("")
    out.write_text("\n".join(lines), encoding="utf-8")
    return out


def _process_row(instance, row: dict[str, str]) -> dict[str, str]:
    workspace_root = instance.root("workspace")
    text_dir = instance.artifact("text_dir")
    path = workspace_root / row["original_path"]
    ext = row.get("extension", "").lower()
    doc_id = row["doc_id"]
    row.setdefault("notes", "")

    if not path.exists():
        row["status_ocr"] = "SOURCE_MISSING"
        row["notes"] = append_note(row.get("notes", ""), "Fichier source introuvable.")
        return row

    try:
        if ext == "pdf":
            pages, page_count, extraction_level = _extract_pdf(path)
        elif ext in TEXT_EXTENSIONS:
            pages, page_count = _extract_plain_text(path)
            extraction_level = "L1_NATIVE_TEXT"
        elif ext in HTML_EXTENSIONS:
            pages, page_count = _extract_html(path)
            extraction_level = "L1_NATIVE_HTML"
        elif ext == "docx":
            pages, page_count = _extract_docx(path)
            extraction_level = "L1_NATIVE_DOCX"
        elif ext == "xlsx":
            pages, page_count = _extract_xlsx(path)
            extraction_level = "L1_NATIVE_XLSX"
        else:
            row["status_ocr"] = "OCR_REQUIRED" if ext in {"png", "jpg", "jpeg", "tif", "tiff", "bmp"} else "UNSUPPORTED"
            row["extraction_level"] = "L2_LOCAL_OCR_REQUIRED" if row["status_ocr"] == "OCR_REQUIRED" else ""
            row["text_quality"] = "missing" if row["status_ocr"] == "OCR_REQUIRED" else "unsupported"
            return row
    except Exception as exc:  # noqa: BLE001
        row["status_ocr"] = "EXTRACTION_ERROR"
        row["notes"] = append_note(row.get("notes", ""), f"Extraction impossible: {exc}")
        return row

    total_chars = sum(len(page.strip()) for page in pages)
    out = _write_page_text(text_dir, doc_id, pages)
    row["text_path"] = relative_to(workspace_root, out)
    row["page_count"] = page_count
    row["text_char_count"] = str(total_chars)
    row["extraction_level"] = extraction_level
    row["text_quality"] = "strong" if total_chars >= 80 else "weak"
    if total_chars >= 80 or ext != "pdf":
        row["status_ocr"] = "TEXT_EXTRACTED"
        if total_chars < 80:
            row["notes"] = append_note(row.get("notes", ""), "Texte court extrait; OCR non applicable a ce format.")
    else:
        row["status_ocr"] = "OCR_REQUIRED"
        row["notes"] = append_note(row.get("notes", ""), "Texte natif faible ou absent: OCR scan requis.")
    return row


def extract_text(instance, run: RunContext, doc_id: str | None = None, docai_mode: str | None = None) -> Path:
    registry_path = instance.register("documents")
    fields, rows = _read_rows(registry_path)
    processed = 0
    for row in rows:
        if doc_id and row.get("doc_id") != doc_id:
            continue
        if row.get("status_ocr") == "OCR_DONE":
            continue
        row.update(_process_row(instance, row))
        text = _text_sample(instance, row) if row.get("text_path") else ""
        apply_access_policy(row, text=text, instance=instance)
        processed += 1
    _write_rows(registry_path, fields, rows)
    run.log_action("write", registry_path, f"text extraction processed={processed}")
    if docai_mode and docai_mode != "off":
        from . import docai

        docai.process_after_native_extract(instance, run, doc_id=doc_id, mode=docai_mode)
    return registry_path


def _load_taxonomy(instance) -> dict:
    return load_structured_file(instance.classification_rules_path())


def _text_sample(instance, row: dict[str, str]) -> str:
    workspace_root = instance.root("workspace")
    parts = [row.get("file_name", ""), row.get("original_path", "")]
    text_path = row.get("text_path")
    if text_path:
        path = workspace_root / text_path
        if path.exists():
            parts.append(path.read_text(encoding="utf-8", errors="ignore")[:6000])
    docling_path = row.get("docling_path")
    if docling_path:
        path = workspace_root / docling_path
        if path.exists():
            parts.append(path.read_text(encoding="utf-8", errors="ignore")[:6000])
    return "\n".join(parts).lower()


def _classify(sample: str, file_name: str, original_path: str, rules: list[dict]) -> tuple[str, str, int]:
    best = ("A_CLASSER", "A_CLASSER", 0)
    file_text = file_name.lower()
    path_text = original_path.lower().replace("\\", "/")
    for rule in rules:
        score = 0
        matched = False
        filename_patterns = rule.get("filename_patterns", rule.get("motifs_nom_fichier", []))
        path_patterns = rule.get("path_patterns", rule.get("motifs_chemin", []))
        keywords = rule.get("keywords", rule.get("mots_cles", []))
        filename_weight = int(rule.get("filename_weight", rule.get("poids_nom_fichier", 50)))
        path_weight = int(rule.get("path_weight", rule.get("poids_chemin", 80)))
        keyword_weight = int(rule.get("keyword_weight", rule.get("poids_mot_cle", 5)))
        document_type = str(rule.get("document_type", rule.get("type_document", "A_CLASSER")))
        for pattern in filename_patterns:
            if _pattern_matches(pattern, file_text):
                matched = True
                score += filename_weight
        for pattern in path_patterns:
            if _pattern_matches(pattern, path_text):
                matched = True
                score += path_weight
        for keyword in keywords:
            if _keyword_matches(sample, str(keyword)):
                matched = True
                score += keyword_weight
        if matched:
            score += int(rule.get("priority", 0))
        if score > best[2]:
            best = (str(rule["lot"]), document_type, score)
    return best


def _normalized_search_text(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", (value or "").lower()).strip()


def _keyword_matches(text: str, keyword: str) -> bool:
    normalized = _normalized_search_text(keyword or "")
    if not normalized:
        return False
    haystack = _normalized_search_text(text)
    parts = [re.escape(part) for part in normalized.split() if part]
    if not parts:
        return False
    pattern = r"(?<![a-z0-9])" + r"\s+".join(parts) + r"(?![a-z0-9])"
    return re.search(pattern, haystack, flags=re.IGNORECASE) is not None


def _pattern_matches(pattern: object, text: str) -> bool:
    raw = str(pattern or "").strip()
    if not raw:
        return False
    if re.fullmatch(r"[A-Za-z0-9]{1,3}", raw):
        return _keyword_matches(text, raw)
    return re.search(raw, text, flags=re.IGNORECASE) is not None


def classify(instance, run: RunContext, copy_files: bool = True) -> Path:
    registry_path = instance.register("documents")
    fields, rows = _read_rows(registry_path)
    taxonomy = _load_taxonomy(instance)
    rules = taxonomy.get("rules", taxonomy.get("regles", []))
    classified_dir = instance.artifact("classified_dir")
    workspace_root = instance.root("workspace")

    processed = 0
    for row in rows:
        sample = _text_sample(instance, row)
        lot, doc_type, score = _classify(sample, row.get("file_name", ""), row.get("original_path", ""), rules)
        row["lot"] = lot
        row["document_type"] = doc_type
        row["suspected_date"] = row.get("suspected_date") or detect_date(sample or row.get("file_name", ""))
        row["classification_status"] = "AUTO_CLASSIFIED" if score else "A_CLASSER"
        apply_access_policy(row, text=sample, instance=instance)
        if copy_files:
            source = workspace_root / row["original_path"]
            target_dir = classified_dir / lot
            target_dir.mkdir(parents=True, exist_ok=True)
            target = target_dir / f"{row['doc_id']}__{safe_name(row['file_name'])}"
            if source.exists() and not target.exists():
                shutil.copy2(source, target)
                run.log_action("write", target, f"classified copy from {source}")
        processed += 1

    write_csv(registry_path, fields, rows)
    run.log_action("write", registry_path, f"classification processed={processed}")
    return registry_path


def _load_proofs(instance) -> list[dict[str, str]]:
    proofs_path = instance.matrix("proofs")
    if proofs_path and proofs_path.exists():
        _, rows = read_csv(proofs_path)
        return rows
    data = load_structured_file(instance.completeness_rules_path())
    return list(data.get("proofs", []))


def _proof_value(proof: dict[str, str], *names: str) -> str:
    aliases = {
        "proof_id": ("proof_id", "preuve_id", "id"),
        "lot": ("lot", "scope", "perimetre"),
        "expected_label": ("expected_label", "libelle_attendu", "preuve_attendue", "document_attendu"),
        "document_type": ("document_type", "type_document", "document_source"),
        "criticality": ("criticality", "criticite", "priorite", "priority"),
        "freshness_months": ("freshness_months", "validity_months", "max_age_months", "fraicheur_mois", "validite_mois"),
    }
    candidates: list[str] = []
    for name in names:
        candidates.extend(aliases.get(name, (name,)))
    for candidate in candidates:
        value = proof.get(candidate)
        if value not in (None, ""):
            return str(value)
    return ""


def _parse_partial_date(value: str) -> date | None:
    match = re.search(r"(20\d{2})(?:[-_/ .]([01]\d)(?:[-_/ .]([0-3]\d))?)?", value or "")
    if not match:
        return None
    year, month, day = match.groups()
    try:
        return date(int(year), int(month or "1"), int(day or "1"))
    except ValueError:
        return None


def _months_between(older: date, newer: date) -> int:
    months = (newer.year - older.year) * 12 + newer.month - older.month
    if newer.day < older.day:
        months -= 1
    return max(0, months)


def _freshness_months(proof: dict[str, str]) -> int | None:
    raw = _proof_value(proof, "freshness_months")
    if not raw:
        return None
    try:
        value = int(raw)
    except ValueError:
        return None
    return value if value > 0 else None


def _tokenize_for_doubt(*values: str) -> set[str]:
    ignored = {"doc", "document", "documents", "piece", "pieces", "attendu", "attendue", "dossier"}
    tokens: set[str] = set()
    for value in values:
        normalized = re.sub(r"[^a-z0-9]+", " ", value.lower())
        tokens.update(token for token in normalized.split() if len(token) >= 4 and token not in ignored)
    return tokens


def _classification_doubts(docs: list[dict[str, str]], lot: str, document_type: str, expected_label: str) -> list[dict[str, str]]:
    tokens = _tokenize_for_doubt(lot, document_type, expected_label)
    doubtful_statuses = {"", "PENDING", "A_CLASSER"}
    candidates: list[dict[str, str]] = []
    seen: set[str] = set()
    for doc in docs:
        if doc.get("doc_id") in seen:
            continue
        same_lot = bool(lot and doc.get("lot") == lot)
        haystack = " ".join([doc.get("file_name", ""), doc.get("original_path", ""), doc.get("notes", "")]).lower()
        token_hit = bool(tokens and any(token in haystack for token in tokens))
        uncertain = doc.get("classification_status", "") in doubtful_statuses or doc.get("document_type", "") in {"", "A_CLASSER"}
        if (same_lot or token_hit) and uncertain:
            candidates.append(doc)
            seen.add(doc.get("doc_id", ""))
    return candidates


def _evidence_paths(matches: list[dict[str, str]]) -> str:
    values: list[str] = []
    for match in matches:
        path = match.get("original_path") or match.get("text_path") or match.get("file_name", "")
        if path:
            values.append(f"{match.get('doc_id', '')}:{path}")
    return "; ".join(values)


def _matched_doc_ids(matches: list[dict[str, str]]) -> str:
    return "; ".join(match.get("doc_id", "") for match in matches if match.get("doc_id"))


def _md_cell(value: str) -> str:
    return (value or "").replace("|", "\\|").replace("\n", " ")


def _status_action(status: str, expected_label: str) -> tuple[str, str]:
    if status == "PRESENT":
        return "Preuve presente", "Aucune demande documentaire immediate."
    if status == "OBSOLETE":
        return (
            f"Demander une version recente: {expected_label}",
            "Demander au syndic une version recente, puis conserver l'ancienne comme preuve historique.",
        )
    if status == "A_CLASSER":
        return (
            f"Verifier le classement avant demande: {expected_label}",
            "Verifier la piece candidate; si elle ne couvre pas l'attendu, demander la piece au syndic.",
        )
    return f"Demander au syndic: {expected_label}", "Demander la piece au syndic et rattacher sa reponse au registre."


def _build_completeness_rows(docs: list[dict[str, str]], proofs: list[dict[str, str]]) -> tuple[list[dict[str, str]], list[dict[str, str]]]:
    by_lot: dict[str, list[dict[str, str]]] = defaultdict(list)
    by_type: dict[str, list[dict[str, str]]] = defaultdict(list)
    for doc in docs:
        by_lot[doc.get("lot", "")].append(doc)
        by_type[doc.get("document_type", "")].append(doc)

    today = datetime.now(timezone.utc).date()
    matrix_rows: list[dict[str, str]] = []
    action_rows: list[dict[str, str]] = []
    for index, proof in enumerate(proofs, start=1):
        proof_id = _proof_value(proof, "proof_id") or f"PRV-{index:03d}"
        lot = _proof_value(proof, "lot")
        expected_label = _proof_value(proof, "expected_label") or proof_id
        doc_type = _proof_value(proof, "document_type")
        criticality = _proof_value(proof, "criticality") or "P2"
        freshness = _freshness_months(proof)
        matches = by_type.get(doc_type, []) if doc_type else by_lot.get(lot, [])
        dated_matches = [(parsed, match) for match in matches if (parsed := _parse_partial_date(match.get("suspected_date", "")))]
        newest_date = max((parsed for parsed, _ in dated_matches), default=None)

        if matches:
            if freshness and newest_date and _months_between(newest_date, today) > freshness:
                status = "OBSOLETE"
            else:
                status = "PRESENT"
            reason = ""
            if status == "OBSOLETE":
                reason = f"Derniere piece datee {newest_date.isoformat()} au-dela du seuil {freshness} mois."
            else:
                reason = "Piece presente dans le registre documentaire."
        else:
            doubt_matches = _classification_doubts(docs, lot, doc_type, expected_label)
            if doubt_matches:
                matches = doubt_matches
                status = "A_CLASSER"
                reason = "Piece candidate trouvee, mais classement documentaire insuffisant ou incertain."
            else:
                status = "ABSENT"
                reason = "Aucune piece locale ne correspond au type attendu."

        subject, action = _status_action(status, expected_label)
        matrix_row = {
            "proof_id": proof_id,
            "lot": lot,
            "expected_label": expected_label,
            "document_type": doc_type,
            "status": status,
            "criticality": criticality,
            "freshness_months": str(freshness or ""),
            "matched_doc_ids": _matched_doc_ids(matches),
            "evidence_paths": _evidence_paths(matches),
            "newest_date": newest_date.isoformat() if newest_date else "",
            "reason": reason,
            "action": action,
        }
        matrix_rows.append(matrix_row)
        if status != "PRESENT":
            action_rows.append(
                {
                    "request_id": f"REQ-DOC-{proof_id}",
                    "source_ref": proof_id,
                    "priority": criticality,
                    "status": status,
                    "subject": subject,
                    "expected_piece": expected_label,
                    "reason": reason,
                    "related_doc_ids": matrix_row["matched_doc_ids"],
                    "evidence_paths": matrix_row["evidence_paths"],
                    "suggested_diligence": action,
                }
            )
    return matrix_rows, action_rows


def _render_completeness_report(instance, docs: list[dict[str, str]], matrix_rows: list[dict[str, str]], action_rows: list[dict[str, str]]) -> str:
    status_counts: dict[str, int] = defaultdict(int)
    for row in matrix_rows:
        status_counts[row["status"]] += 1

    lines = [
        "# Rapport de completude documentaire",
        "",
        f"- Instance: {instance.display_name}",
        f"- Documents inventories: {len(docs)}",
        f"- Pieces attendues: {len(matrix_rows)}",
        f"- Pieces a traiter: {len(action_rows)}",
        "",
        "## Synthese",
        "",
        "| Statut | Nombre | Lecture CS |",
        "|---|---:|---|",
        f"| PRESENT | {status_counts['PRESENT']} | Piece locale suffisante. |",
        f"| ABSENT | {status_counts['ABSENT']} | Piece a demander au syndic. |",
        f"| OBSOLETE | {status_counts['OBSOLETE']} | Version recente a obtenir. |",
        f"| A_CLASSER | {status_counts['A_CLASSER']} | Classement a verifier avant relance. |",
        "",
        "## Matrice actionnable",
        "",
        "| Statut | Priorite | Lot | Piece attendue | Type source | Preuves | Suite |",
        "|---|---|---|---|---|---|---|",
    ]
    for row in matrix_rows:
        evidence = row["evidence_paths"] or row["matched_doc_ids"]
        lines.append(
            f"| {_md_cell(row['status'])} | {_md_cell(row['criticality'])} | {_md_cell(row['lot'])} | "
            f"{_md_cell(row['expected_label'])} | {_md_cell(row['document_type'])} | "
            f"{_md_cell(evidence)} | {_md_cell(row['action'])} |"
        )

    lines.extend(
        [
            "",
            "## Pieces a demander",
            "",
            "| Priorite | Statut | Piece | Pourquoi | Diligence proposee |",
            "|---|---|---|---|---|",
        ]
    )
    if action_rows:
        for row in action_rows:
            lines.append(
                f"| {_md_cell(row['priority'])} | {_md_cell(row['status'])} | {_md_cell(row['expected_piece'])} | "
                f"{_md_cell(row['reason'])} | {_md_cell(row['suggested_diligence'])} |"
            )
    else:
        lines.append("| OK | PRESENT | Aucune | Toutes les pieces attendues sont couvertes. | Aucune relance. |")
    return "\n".join(lines) + "\n"


def missing_docs(instance, run: RunContext) -> Path:
    _, docs = read_csv(instance.register("documents"))
    proofs = _load_proofs(instance)
    reports_dir = instance.artifact("reports_dir")
    findings_path = instance.register("findings")
    _, existing_findings = read_csv(findings_path)
    report_path = reports_dir / "rapport_completude_documentaire.md"
    matrix_path = reports_dir / "matrice_completude_documentaire.csv"
    requests_path = reports_dir / "pieces_a_demander.csv"
    matrix_rows, action_rows = _build_completeness_rows(docs, proofs)
    finding_rows: list[dict[str, str]] = []
    for row in matrix_rows:
        if row["status"] != "PRESENT":
            finding_rows.append(
                {
                    "finding_id": f"FDG-{row['proof_id']}",
                    "category": "completeness",
                    "severity": row["criticality"],
                    "source_ref": row["proof_id"],
                    "fact": f"{row['status']}: {row['expected_label']}",
                    "diligence": row["action"],
                    "status": "OPEN",
                }
            )

    lines = _render_completeness_report(instance, docs, matrix_rows, action_rows).splitlines()
    extranet_path = instance.matrix("extranet")
    if extranet_path and extranet_path.exists():
        _, extranet_rows = read_csv(extranet_path)
        lines.extend(
            [
                "",
                "## Matrice extranet reference",
                "",
                "| Scope | Document attendu | Present extranet | Qualification |",
                "|---|---|---|---|",
            ]
        )
        for row in extranet_rows[:30]:
            lines.append(
                f"| {row.get('scope', '')} | {row.get('document_attendu', '')} | "
                f"{row.get('present_extranet', '') or 'A_VERIFIER'} | {row.get('qualification', '')} |"
            )

    write_csv(matrix_path, COMPLETENESS_FIELDS, matrix_rows)
    write_csv(requests_path, DOCUMENT_REQUEST_FIELDS, action_rows)
    write_text(report_path, "\n".join(lines) + "\n")
    merged = [row for row in existing_findings if row.get("category") != "completeness"] + finding_rows
    write_csv(findings_path, FINDING_FIELDS, merged)
    run.log_action("write", matrix_path, f"document completeness rows={len(matrix_rows)}")
    run.log_action("write", requests_path, f"document requests rows={len(action_rows)}")
    run.log_action("write", report_path, f"missing docs report rows={len(proofs)}")
    run.log_action("write", findings_path, f"findings merged={len(merged)}")
    return report_path


def _pct(part: int, total: int) -> str:
    if total == 0:
        return "0.0%"
    return f"{(part / total) * 100:.1f}%"


def compute_kpis(instance, run: RunContext) -> Path:
    kpi_path = instance.register("kpi")
    if kpi_path.exists():
        fields, rows = read_csv(kpi_path)
    else:
        template = instance.register("kpi")
        template.parent.mkdir(parents=True, exist_ok=True)
        template.write_text(load_template_csv("kpi.csv"), encoding="utf-8")
        fields, rows = read_csv(template)
        run.log_action("write", template, "bootstrap kpi register")

    _, docs = read_csv(instance.register("documents"))
    total_docs = len(docs)
    ocr_required = sum(1 for doc in docs if doc.get("status_ocr") == "OCR_REQUIRED")
    unclassified = sum(1 for doc in docs if doc.get("classification_status") in {"A_CLASSER", "PENDING", ""})
    request_stats = request_metrics(instance)

    values = {
        "KPI-001": _pct(ocr_required, total_docs),
        "KPI-002": _pct(unclassified, total_docs),
        "KPI-011": _pct(request_stats["traced"], request_stats["total"]),
        "KPI-012": _pct(request_stats["complete"], request_stats["total"]),
    }
    for row in rows:
        kpi_id = row.get("kpi_id", "")
        if kpi_id in values:
            row["value"] = values[kpi_id]
            row["valeur"] = values[kpi_id]
            row["status"] = "CALCULATED"
            row["statut"] = "CALCULE"
    write_csv(kpi_path, list(fields or rows[0].keys() if rows else []), rows)
    run.log_action("write", kpi_path, f"kpi updated ids={','.join(values)}")
    return kpi_path
